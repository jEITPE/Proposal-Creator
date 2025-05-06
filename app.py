import os
import re
import base64
import uuid
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from flask import Flask, render_template, request, redirect, url_for, send_file, abort, jsonify, session, flash
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from bs4 import BeautifulSoup
import json
from functools import wraps
import datetime
import shutil
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import html
import logging
from dotenv import load_dotenv
from models import db, Usuario, Perfil, BlocoProposta, Oferta, BlocoPropostaOferta, Proposta, Rascunho, UsuarioBloco
from docxtpl import DocxTemplate, InlineImage
from datetime import timedelta
import requests
import sys
import time
import random
import string
import subprocess
from docxcompose.composer import Composer
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.interval import IntervalTrigger
import atexit
from migrate_data import migrar_blocos, migrar_ofertas, migrar_propostas, migrar_rascunhos, sincronizar_banco_para_json
from migrate_users import migrar_usuarios_do_json
from db_operations import (
    obter_usuarios_db, salvar_usuario_db,
    obter_blocos_db, salvar_bloco_db,
    obter_ofertas_db, salvar_oferta_db,
    obter_propostas_db, salvar_proposta_db, criar_proposta_db,
    obter_rascunhos_db, salvar_rascunho_db
)

# Configurações de logging
LOGS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
os.makedirs(LOGS_FOLDER, exist_ok=True)

# Configurar logging para arquivo e console
log_file = os.path.join(LOGS_FOLDER, 'app.log')
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)

# Manter referência ao logger da aplicação
logger = logging.getLogger(__name__)

# Definir cores da Service IT
SERVICE_IT_RED = RGBColor(230, 0, 0)  # Vermelho principal
SERVICE_IT_DARK_RED = RGBColor(180, 0, 0)  # Vermelho escuro para destaques
SERVICE_IT_BLACK = RGBColor(0, 0, 0)  # Preto para textos principais
SERVICE_IT_GRAY = RGBColor(128, 128, 128)  # Cinza para textos secundários

# Função para verificar se uma fonte está disponível e retornar a fonte ou um fallback
def get_font_name(primary_font, fallback_font='Calibri'):
    try:
        # Lista de caminhos comuns para fontes
        font_paths = [
            os.path.join(app.root_path, 'static', 'fonts'),  # Pasta de fontes no projeto
            '/usr/share/fonts',                             # Pasta de fontes no Linux
            '/usr/local/share/fonts',                       # Outra pasta no Linux
            'C:\\Windows\\Fonts',                          # Pasta de fontes no Windows
            os.path.expanduser('~/Library/Fonts'),         # Pasta de fontes no MacOS
            os.path.expanduser('~/.fonts')                 # Pasta de usuário no Linux
        ]
        
        # Se for Poppins, retornar o nome mesmo que não tenha sido encontrada
        if primary_font.lower() == 'poppins':
            return primary_font
            
        # Verificar se a fonte existe em algum dos caminhos
        for font_path in font_paths:
            if os.path.exists(font_path):
                for root, dirs, files in os.walk(font_path):
                    for file in files:
                        if file.lower().endswith(('.ttf', '.otf')) and primary_font.lower() in file.lower():
                            return primary_font
                    
        # Se chegou aqui, não encontrou a fonte, retornar o fallback
        logging.warning(f"Fonte {primary_font} não encontrada, usando {fallback_font}")
        return fallback_font
    except Exception as e:
        # Fallback seguro em caso de qualquer erro
        logging.warning(f"Erro ao verificar fonte {primary_font}: {e}, usando {fallback_font}")
        return fallback_font

# Configurar aplicação Flask
app = Flask(__name__, static_url_path='/static')
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'chave-secreta-padrao')
app.config['UPLOAD_FOLDER'] = os.path.join('static', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB (limite de tamanho de arquivo)

# Configuração do banco de dados
if os.environ.get('DATABASE_HOST'):
    # Configuração para PostgreSQL
    db_user = os.environ.get('DATABASE_USER', 'postgres')
    db_password = os.environ.get('DATABASE_PASSWORD', 'postgres')
    db_host = os.environ.get('DATABASE_HOST', 'localhost')
    db_name = os.environ.get('DATABASE_NAME', 'proposal_creator')
    db_port = os.environ.get('DATABASE_PORT', '5432')
    
    # URL de conexão com opcões explícitas de codificação
    db_url = f"postgresql://{db_user}:{db_password}@{db_host}:{db_port}/{db_name}"
    
    app.config['SQLALCHEMY_DATABASE_URI'] = db_url
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    
    # Configurações de engine para lidar corretamente com codificação UTF-8
    app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
        'connect_args': {
            'client_encoding': 'UTF8',
            'options': '-c client_encoding=UTF8 -c standard_conforming_strings=on'
        },
        'echo': True,  # Para debug
        'isolation_level': 'READ COMMITTED'
    }
else:
    # Configuração para SQLite (caso não tenha configuração de PostgreSQL)
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///proposals.db'
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Inicializar extensões
db.init_app(app)

# Garantir que os diretórios necessários existam
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join('static', 'img'), exist_ok=True)
os.makedirs(os.path.join('data'), exist_ok=True)
os.makedirs(os.path.join(app.root_path, 'static', 'fonts'), exist_ok=True)

# Função para garantir que a fonte Poppins esteja disponível
def garantir_fonte_poppins():
    try:
        # Caminhos para os arquivos de fonte
        fonts_dir = os.path.join(app.root_path, 'static', 'fonts')
        os.makedirs(fonts_dir, exist_ok=True)
        
        poppins_regular = os.path.join(fonts_dir, 'Poppins-Regular.ttf')
        poppins_bold = os.path.join(fonts_dir, 'Poppins-Bold.ttf')
        poppins_italic = os.path.join(fonts_dir, 'Poppins-Italic.ttf')
        
        # Verificar se as fontes já existem e logar informação
        if os.path.exists(poppins_regular):
            logging.info(f"Fonte Poppins Regular já existe em: {poppins_regular}")
        else:
            logging.warning(f"Fonte Poppins Regular não encontrada em: {poppins_regular}")
            
        if os.path.exists(poppins_bold):
            logging.info(f"Fonte Poppins Bold já existe em: {poppins_bold}")
        else:
            logging.warning(f"Fonte Poppins Bold não encontrada em: {poppins_bold}")
            
        if os.path.exists(poppins_italic):
            logging.info(f"Fonte Poppins Italic já existe em: {poppins_italic}")
        else:
            logging.warning(f"Fonte Poppins Italic não encontrada em: {poppins_italic}")
        
        # Se todas as fontes existirem, retornar sucesso
        if os.path.exists(poppins_regular) and os.path.exists(poppins_bold) and os.path.exists(poppins_italic):
            logging.info("Todas as variantes da fonte Poppins já existem no sistema.")
            return True
        
        # URLs para download das fontes Poppins
        poppins_urls = {
            'regular': "https://github.com/google/fonts/raw/main/ofl/poppins/Poppins-Regular.ttf",
            'bold': "https://github.com/google/fonts/raw/main/ofl/poppins/Poppins-Bold.ttf",
            'italic': "https://github.com/google/fonts/raw/main/ofl/poppins/Poppins-Italic.ttf"
        }
        
        # Baixar as fontes que faltam
        for variant, url in poppins_urls.items():
            dest_path = None
            if variant == 'regular' and not os.path.exists(poppins_regular):
                dest_path = poppins_regular
            elif variant == 'bold' and not os.path.exists(poppins_bold):
                dest_path = poppins_bold
            elif variant == 'italic' and not os.path.exists(poppins_italic):
                dest_path = poppins_italic
                
            if dest_path:
                try:
                    logging.info(f"Baixando fonte Poppins {variant} de {url}...")
                    response = requests.get(url, stream=True, timeout=15)
                    if response.status_code == 200:
                        with open(dest_path, 'wb') as f:
                            for chunk in response.iter_content(chunk_size=8192):
                                if chunk:
                                    f.write(chunk)
                        logging.info(f"Fonte Poppins {variant} baixada com sucesso para {dest_path}")
                    else:
                        logging.error(f"Falha ao baixar fonte Poppins {variant}. Status code: {response.status_code}")
                except Exception as e:
                    logging.error(f"Erro ao baixar fonte Poppins {variant}: {e}")
        
        # Verificar novamente se todas as fontes foram baixadas
        if os.path.exists(poppins_regular) and os.path.exists(poppins_bold) and os.path.exists(poppins_italic):
            logging.info("Todas as fontes Poppins foram baixadas com sucesso!")
            return True
        else:
            # Se pelo menos a Regular existir, podemos prosseguir
            if os.path.exists(poppins_regular):
                logging.info("Pelo menos a fonte Poppins Regular está disponível, continuando...")
                return True
            else:
                logging.error("Não foi possível baixar a fonte Poppins Regular, verifique a conexão com a internet.")
            return False
    except Exception as e:
        logging.error(f"Erro ao garantir fonte Poppins: {e}")
        # Mesmo com erro, tentamos continuar se pelo menos a Regular existir
        if os.path.exists(os.path.join(fonts_dir, 'Poppins-Regular.ttf')):
            logging.info("Apesar do erro, a fonte Poppins Regular existe, continuando...")
            return True
        return False

# Garantir que a fonte Poppins esteja disponível
garantir_fonte_poppins()

# Inicializar o modelo de banco de dados
try:
    from models import db
    db.init_app(app)
    with app.app_context():
        # Verificar e configurar codificação do cliente
        from sqlalchemy import text
        result = db.session.execute(text("SHOW client_encoding"))
        client_encoding = result.scalar()
        app.logger.info(f"Codificação do cliente: {client_encoding}")
        
        if client_encoding.lower() != 'utf8':
            app.logger.warning(f"Ajustando client_encoding para UTF8 (era {client_encoding})...")
            db.session.execute(text("SET client_encoding TO 'UTF8'"))
        
        # Verificar codificação do servidor
        result = db.session.execute(text("SHOW server_encoding"))
        server_encoding = result.scalar()
        app.logger.info(f"Codificação do servidor: {server_encoding}")
        
        # Criar tabelas se não existirem
        db.create_all()
        app.logger.info("Tabelas criadas ou verificadas com sucesso!")
except Exception as e:
    app.logger.error(f"Erro ao inicializar tabelas: {e}")
    # Modo fallback para SQLite em caso de problemas com PostgreSQL
    # app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///proposals.db'

# Garantir que todas as consultas ao banco de dados usem o contexto de aplicação
def with_app_context(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        with app.app_context():
            return func(*args, **kwargs)
    return wrapper


# Arquivo para armazenar usuários
USUARIOS_FILE = os.path.join('data', 'usuarios.json')

# Arquivo para armazenar ofertas
OFERTAS_FILE = os.path.join('data', 'ofertas.json')

# Extensões permitidas para upload
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Função para carregar usuários - Modificada para banco de dados
def carregar_usuarios():
    """
    Carrega os usuários diretamente do banco de dados.
    Retorna um dicionário de usuários.
    """
    try:
        return obter_usuarios_db()
    except Exception as e:
        logger.error(f"Erro ao carregar usuários do banco: {e}")
        return {}

# Função para salvar usuários - Modificada para banco de dados
def salvar_usuarios(usuarios):
    """
    Salva usuários diretamente no banco de dados.
    """
    try:
        for login, dados in usuarios.items():
            salvar_usuario_db(login, dados)
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar usuários no banco: {e}")
        return False

# Arquivo para armazenar propostas
PROPOSTAS_FILE = os.path.join('data', 'propostas.json')

# Função para carregar propostas - Modificada para banco de dados
def carregar_propostas():
    """
    Carrega propostas diretamente do banco de dados.
    """
    try:
        return obter_propostas_db()
    except Exception as e:
        logger.error(f"Erro ao carregar propostas do banco: {e}")
        return {}

# Função para salvar propostas - Modificada para banco de dados
def salvar_propostas(propostas):
    """
    Salva propostas diretamente no banco de dados.
    """
    try:
        for id_proposta, dados in propostas.items():
            salvar_proposta_db(id_proposta, dados)
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar propostas no banco: {e}")
        return False

# Arquivo para armazenar blocos
BLOCOS_FILE = os.path.join('data', 'blocos.json')

# Função para carregar blocos - Modificada para banco de dados
def carregar_blocos():
    """
    Carrega blocos diretamente do banco de dados.
    """
    try:
        return obter_blocos_db()
    except Exception as e:
        logger.error(f"Erro ao carregar blocos do banco: {e}")
        return {}

# Função para salvar blocos - Modificada para banco de dados
def salvar_blocos(blocos):
    """
    Salva blocos diretamente no banco de dados.
    """
    try:
        for nome_bloco, dados in blocos.items():
            salvar_bloco_db(nome_bloco, dados)
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar blocos no banco: {e}")
        return False

# Função para carregar rascunhos - Modificada para banco de dados
def carregar_rascunhos():
    """
    Carrega rascunhos diretamente do banco de dados.
    """
    try:
        return obter_rascunhos_db()
    except Exception as e:
        logger.error(f"Erro ao carregar rascunhos do banco: {e}")
        return {}

# Função para salvar rascunhos - Modificada para banco de dados
def salvar_rascunhos(rascunhos):
    """
    Salva rascunhos diretamente no banco de dados.
    """
    try:
        for id_rascunho, dados in rascunhos.items():
            salvar_rascunho_db(id_rascunho, dados)
        return True
    except Exception as e:
        logger.error(f"Erro ao salvar rascunhos no banco: {e}")
        return False

# Decorator para verificar se o usuário está logado
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'usuario_logado' not in session:
            flash('Por favor, faça login para acessar esta página.', 'warning')
            return redirect(url_for('login', next=request.url))
        
        try:
            # Verificar se o usuário existe e está ativo
            with app.app_context():
                usuario = Usuario.query.filter_by(login=session['usuario_logado']).first()
                if not usuario:
                    flash('Usuário não encontrado. Por favor, faça login novamente.', 'warning')
                    return redirect(url_for('logout'))
                    
                # Verificar se o usuário está ativo
                if usuario.status != 1:
                    flash('Sua conta está desativada. Entre em contato com o administrador.', 'warning')
                    return redirect(url_for('logout'))
                    
                # Verificar se o acesso do usuário temporário expirou
                if usuario.data_expiracao and datetime.datetime.utcnow() > usuario.data_expiracao:
                    flash('Seu acesso temporário expirou. Entre em contato com o administrador.', 'warning')
                    return redirect(url_for('logout'))
                    
                # Atualizar último acesso
                usuario.ultimo_acesso = datetime.datetime.utcnow()
                db.session.commit()
                
                return f(*args, **kwargs)
        except Exception as e:
            app.logger.error(f"Erro ao verificar permissões do usuário: {e}")
            
            # Fallback para o método antigo em caso de erro
            usuarios = carregar_usuarios()
            if session['usuario_logado'] not in usuarios:
                flash('Sessão inválida. Por favor, faça login novamente.', 'warning')
                return redirect(url_for('logout'))
            
            # Verificar status do usuário no JSON (fallback)
            if 'status' in usuarios[session['usuario_logado']] and usuarios[session['usuario_logado']]['status'] != 1:
                flash('Sua conta está desativada. Entre em contato com o administrador.', 'warning')
                return redirect(url_for('logout'))
        
        return f(*args, **kwargs)
    return decorated_function

# Decorator para verificar se o usuário é admin
def admin_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        # Verificar se o usuário está logado e é administrador
        if 'usuario_logado' not in session:
            flash('Por favor, faça login para acessar esta página.', 'warning')
            return redirect(url_for('login'))
        
        try:
            # Primeiro tenta verificar no banco de dados
            with app.app_context():
                usuario = Usuario.query.filter_by(login=session['usuario_logado']).first()
                if usuario:
                    perfil = Perfil.query.get(usuario.id_perfil)
                    if perfil and perfil.nome == "Governança":
                        return f(*args, **kwargs)
                    else:
                        flash('Acesso negado. Você precisa ser administrador para acessar esta página.', 'danger')
                        return redirect(url_for('dashboard'))
        except Exception as e:
            app.logger.error(f"Erro ao verificar permissões do usuário: {e}")
            # Fallback para verificar no JSON em caso de erro
            usuarios = carregar_usuarios()
            if session['usuario_logado'] not in usuarios or usuarios[session['usuario_logado']]['tipo'] != 'admin':
                flash('Acesso negado. Você precisa ser administrador para acessar esta página.', 'danger')
                return redirect(url_for('dashboard'))
        
        return f(*args, **kwargs)
    return decorated_function

# Decorator para verificar se o usuário tem permissão para editar um bloco
def permissao_bloco_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        # Obter o nome do bloco a ser editado
        bloco_nome = kwargs.get('bloco_nome') or request.args.get('bloco_nome') or request.form.get('bloco_nome')
        
        if not bloco_nome:
            flash('Bloco não especificado.', 'warning')
            return redirect(url_for('dashboard'))
            
        try:
            # Verificar se o usuário tem permissão para editar o bloco
            with app.app_context():
                usuario = Usuario.query.filter_by(login=session['usuario_logado']).first()
                if not usuario:
                    flash('Usuário não encontrado.', 'danger')
                    return redirect(url_for('logout'))
                
                # Se for admin ou superusuário, tem permissão total
                perfil = Perfil.query.get(usuario.id_perfil)
                if perfil and perfil.nome == "Governança" or usuario.superusuario:
                    return f(*args, **kwargs)
                
                # Verificar nas permissões específicas
                permitido = False
                for bloco in usuario.blocos_permitidos:
                    if bloco.nome == bloco_nome:
                        permitido = True
                        break
                
                if not permitido:
                    flash(f'Você não tem permissão para editar o bloco {bloco_nome}.', 'danger')
                    return redirect(url_for('dashboard'))
                
                return f(*args, **kwargs)
        except Exception as e:
            app.logger.error(f"Erro ao verificar permissão para editar bloco: {e}")
            flash('Erro ao verificar permissões. Contate o administrador.', 'danger')
            return redirect(url_for('dashboard'))
        
        return f(*args, **kwargs)
    return decorated_function

# Rota de login
@app.route('/login', methods=['GET', 'POST'])
def login():
    # Se o usuário já estiver logado, redireciona para o dashboard
    if 'usuario_logado' in session:
        return redirect(url_for('dashboard'))
        
    error = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        try:
            with app.app_context():
                # Busca o usuário no banco de dados
                usuario_db = Usuario.query.filter_by(login=username, status=1).first()
                
                # Verificar se o usuário existe e se a senha está correta
                senha_valida = False
                if usuario_db:
                    # Se a senha começa com o prefixo de hash, usa check_password_hash
                    if usuario_db.senha.startswith('pbkdf2:sha256:'):
                        senha_valida = check_password_hash(usuario_db.senha, password)
                    else:
                        # Compatibilidade com senhas em texto puro
                        senha_valida = usuario_db.senha == password
                        
                        # Se a senha é válida mas não está em hash, atualizar para hash
                        if senha_valida:
                            usuario_db.senha = generate_password_hash(password)
                            db.session.commit()
                            app.logger.info(f"Senha do usuário {username} atualizada para formato hash")
                
                if usuario_db and senha_valida:
                    session['usuario_logado'] = usuario_db.login
                    # Define o tipo de usuário conforme o perfil
                    perfil = Perfil.query.get(usuario_db.id_perfil)
                    if perfil and perfil.nome == "Governança":
                        session['tipo_usuario'] = 'admin'
                    else:
                        session['tipo_usuario'] = 'usuario'
                    
                    # Atualiza o último acesso
                    usuario_db.ultimo_acesso = datetime.datetime.now()
                    db.session.commit()
                    
                    return redirect(url_for('dashboard'))
                else:
                    error = 'Usuário ou senha inválidos. Tente novamente.'
        except Exception as e:
            app.logger.error(f"Erro ao autenticar usuário do banco: {e}")
            # Fallback para JSON em caso de erro no banco de dados
            usuarios = carregar_usuarios()
            if username in usuarios and usuarios[username]['senha'] == password:
                session['usuario_logado'] = username
                session['tipo_usuario'] = usuarios[username]['tipo']
                return redirect(url_for('dashboard'))
            else:
                error = 'Erro de autenticação. Tente novamente.'
    
    return render_template('login.html', error=error)

# Rota de logout
@app.route('/logout')
def logout():
    session.pop('usuario_logado', None)
    session.pop('tipo_usuario', None)
    return redirect(url_for('login'))

# Rota para o dashboard
@app.route('/dashboard', methods=['GET'])
@login_required
def dashboard():
    # Carregar propostas existentes
    propostas = carregar_propostas()
    
    # Obter parâmetros de filtro
    filtro_busca = request.args.get('busca', '')
    filtro_data = request.args.get('data', '')
    ordem = request.args.get('ordem', 'recentes')  # recentes (padrão) ou antigos
    
    # Formatar data para o formato usado nas propostas (dd/mm/aaaa)
    data_formatada = ""
    if filtro_data:
        try:
            # Converter do formato yyyy-mm-dd (HTML input date) para dd/mm/yyyy (formato da aplicação)
            data_obj = datetime.datetime.strptime(filtro_data, '%Y-%m-%d')
            data_formatada = data_obj.strftime('%d/%m/%Y')
        except:
            data_formatada = filtro_data
    
    # Carregar rascunhos do usuário
    rascunhos = carregar_rascunhos()
    rascunhos_usuario = {}
    
    # Filtrar rascunhos do usuário atual
    for rascunho_id, rascunho in rascunhos.items():
        if rascunho.get('usuario') == session.get('usuario_logado'):
            rascunhos_usuario[rascunho_id] = rascunho
    
    # Processar as propostas (para ordenação e filtros)
    propostas_filtradas = {}
    propostas_lista = []
    
    for proposta_id, proposta in propostas.items():
        # Adicionar id à proposta para facilitar ordenação
        proposta_com_id = proposta.copy()
        proposta_com_id['id'] = proposta_id
        
        # Aplicar filtro de busca (cliente ou gerado_por)
        if filtro_busca and filtro_busca.lower() not in proposta.get('nome_cliente', '').lower() and filtro_busca.lower() not in proposta.get('gerado_por', '').lower():
            continue
            
        # Aplicar filtro de data
        if data_formatada and data_formatada not in proposta.get('data_geracao', ''):
            continue
            
        # Adicionar à lista para ordenação
        propostas_lista.append(proposta_com_id)
    
    # Ordenar as propostas por data
    try:
        if ordem == 'recentes':
            propostas_lista.sort(key=lambda x: datetime.datetime.strptime(x.get('data_geracao', '01/01/1970 00:00:00'), '%d/%m/%Y %H:%M:%S'), reverse=True)
        else:
            propostas_lista.sort(key=lambda x: datetime.datetime.strptime(x.get('data_geracao', '01/01/1970 00:00:00'), '%d/%m/%Y %H:%M:%S'))
    except Exception as e:
        logger.error(f"Erro ao ordenar propostas: {e}")
    
    # Converter lista ordenada de volta para dicionário
    for proposta in propostas_lista:
        proposta_id = proposta.pop('id')  # Remover id temporário
        propostas_filtradas[proposta_id] = proposta
    
    # Contar propostas de hoje
    hoje = datetime.datetime.now().strftime("%d/%m/%Y")
    propostas_hoje = 0
    for proposta in propostas.values():
        if hoje in proposta.get('data_geracao', ''):
            propostas_hoje += 1
    
    # Verificar tipo de usuário para permissões
    tipo_usuario = session.get('tipo_usuario', 'usuario')
    
    return render_template('dashboard.html', 
                          propostas=propostas_filtradas, 
                          rascunhos_usuario=rascunhos_usuario, 
                          tipo_usuario=tipo_usuario,
                          propostas_hoje=propostas_hoje,
                          filtro_busca=filtro_busca,
                          filtro_data=filtro_data,
                          ordem=ordem)

# Rota para visualizar uma proposta específica
@app.route('/visualizar_proposta/<proposta_id>')
@login_required
def visualizar_proposta(proposta_id):
    propostas = carregar_propostas()
    if proposta_id in propostas:
        proposta = propostas[proposta_id]
        
        # Carregar blocos específicos para este cliente
        blocos = carregar_blocos()
        blocos_cliente = []
        
        # Filtrar blocos específicos para este cliente
        for nome_bloco, info_bloco in blocos.items():
            if info_bloco.get('cliente_associado') == proposta['nome_cliente']:
                blocos_cliente.append(nome_bloco)
        
        return render_template('visualizar_proposta.html', 
                              proposta=proposta,
                              proposta_id=proposta_id,
                              blocos_cliente=blocos_cliente)
    else:
        flash('Proposta não encontrada.')
        return redirect(url_for('dashboard'))

# Rota para baixar uma proposta
@app.route('/baixar_proposta/<proposta_id>')
@login_required
def baixar_proposta(proposta_id):
    propostas = carregar_propostas()
    if proposta_id in propostas:
        proposta = propostas[proposta_id]
        arquivo = proposta.get('arquivo', '')
        
        logger.info(f"Tentando baixar arquivo: {arquivo}")
        
        # Verificar se o caminho é absoluto ou relativo
        if arquivo:
            # Se o caminho já for absoluto, usar como está
            if os.path.isabs(arquivo):
                caminho_completo = arquivo
            else:
                # Se for relativo, combinar com o caminho raiz da aplicação
                caminho_completo = os.path.join(app.root_path, arquivo)
            
            logger.info(f"Caminho completo do arquivo: {caminho_completo}")
            
            # Verificar se o arquivo existe no caminho completo
            if os.path.exists(caminho_completo):
                logger.info(f"Arquivo encontrado, enviando: {caminho_completo}")
                return send_file(caminho_completo, as_attachment=True)
            else:
                logger.error(f"Arquivo não encontrado no caminho: {caminho_completo}")
                
                # Tentar buscar apenas pelo nome do arquivo no diretório de uploads
                nome_arquivo = os.path.basename(arquivo)
                caminho_alternativo = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
                logger.info(f"Tentando caminho alternativo: {caminho_alternativo}")
                
                if os.path.exists(caminho_alternativo):
                    logger.info(f"Arquivo encontrado no caminho alternativo: {caminho_alternativo}")
                    return send_file(caminho_alternativo, as_attachment=True)
                
                flash('Arquivo da proposta não encontrado.')
        else:
            flash('Caminho do arquivo não definido.')
    else:
        flash('Proposta não encontrada.')
    
    return redirect(url_for('dashboard'))

# Rota para excluir uma proposta
@app.route('/excluir_proposta/<proposta_id>')
@login_required
def excluir_proposta(proposta_id):
    # Carregar propostas existentes
    propostas = carregar_propostas()
    if proposta_id in propostas:
        # Verificar se o usuário é admin ou o criador da proposta
        if session.get('tipo_usuario') == 'admin' or propostas[proposta_id].get('gerado_por') == session.get('usuario_logado'):
            # Remover o arquivo se existir
            arquivo = propostas[proposta_id].get('arquivo', '')
            if arquivo and os.path.exists(arquivo):
                try:
                    os.remove(arquivo)
                except:
                    pass
            
            # Remover a proposta do banco de dados
            try:
                proposta_db = Proposta.query.get(proposta_id)
                if proposta_db:
                    db.session.delete(proposta_db)
                    db.session.commit()
                    app.logger.info(f"Proposta {proposta_id} removida do banco de dados")
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"Erro ao remover proposta do banco de dados: {str(e)}")
            
            # Remover a proposta do dicionário JSON
            del propostas[proposta_id]
            # Salvar as alterações no JSON
            salvar_propostas(propostas)
            
            flash('Proposta excluída com sucesso.')
        else:
            flash('Você não tem permissão para excluir esta proposta.')
    else:
        flash('Proposta não encontrada.')
    return redirect(url_for('dashboard'))

# Rota para criar uma nova proposta (exibir formulário)
@app.route('/criar_proposta', methods=['GET'])
@login_required
def exibir_criar_proposta():
    # Verificar se é uma nova proposta (reiniciar)
    reiniciar = request.args.get('reiniciar', 'false') == 'true'
    
    # Obter o cliente da URL, se fornecido
    cliente = request.args.get('cliente', '')
    
    # Obter informações do usuário atual
    usuario_atual = session.get('usuario_logado')
    tipo_usuario = session.get('tipo_usuario')
    is_admin = tipo_usuario == 'admin'
    
    # Carregar todos os blocos disponíveis
    todos_blocos = carregar_blocos()
    
    # Filtrar blocos com base nas permissões do usuário
    blocos = {}
    
    try:
        # Obter o usuário do banco de dados
        usuario_obj = None
        is_superusuario = False
        
        try:
            usuario_obj = Usuario.query.filter_by(login=usuario_atual).first()
            is_superusuario = usuario_obj and usuario_obj.superusuario
        except Exception as e:
            logger.error(f"Erro ao obter usuário do banco: {e}")
            # Fallback para JSON em caso de erro
            usuarios = carregar_usuarios()
            is_superusuario = usuarios.get(usuario_atual, {}).get('superusuario', False)
        
        if is_admin or is_superusuario:
            # Admins e superusuários têm acesso a todos os blocos
            blocos = todos_blocos
        else:
            # Filtrar blocos para usuários regulares
            if usuario_obj:
                # Obter blocos permitidos do banco de dados
                for bloco_db in BlocoProposta.query.all():
                    tem_permissao = False
                    
                    # Verificar se o bloco é obrigatório (todos têm acesso)
                    if bloco_db.obrigatorio:
                        tem_permissao = True
                    # Verificar se o usuário é o criador do bloco
                    elif bloco_db.criado_por == usuario_atual:
                        tem_permissao = True
                    # Verificar se o usuário tem permissão específica para o bloco
                    elif bloco_db in usuario_obj.blocos_permitidos:
                        tem_permissao = True
                    
                    if tem_permissao and bloco_db.nome in todos_blocos:
                        blocos[bloco_db.nome] = todos_blocos[bloco_db.nome]
            else:
                # Fallback para JSON se não for possível acessar o banco
                # Apenas blocos obrigatórios serão mostrados
                for nome_bloco, bloco in todos_blocos.items():
                    if bloco.get('obrigatorio', False):
                        blocos[nome_bloco] = bloco
    except Exception as e:
        logger.error(f"Erro ao filtrar blocos por permissão: {e}")
        # Em caso de erro, mostrar apenas blocos obrigatórios
        for nome_bloco, bloco in todos_blocos.items():
            if bloco.get('obrigatorio', False):
                blocos[nome_bloco] = bloco
    
    # Carregar ofertas disponíveis
    ofertas = carregar_ofertas()
    
    # Carregar rascunhos
    rascunhos = carregar_rascunhos()
    
    # Verificar se há um rascunho para este cliente
    rascunho = None
    rascunho_id = None
    
    if not reiniciar and cliente:
        # Procurar rascunho para este cliente
        for r_id, r_data in rascunhos.items():
            if r_data.get('nome_cliente') == cliente:
                rascunho = r_data
                rascunho_id = r_id
                break
    
    # Adicionar mensagem informativa sobre o formato das variáveis
    flash('Lembre-se: Use {{NOME_CLIENTE}} para o nome do cliente e {{logo_cliente}} para o logo na capa.')
    
    # Preparar dados para o template
    template_data = {
        'cliente': cliente,
        'blocos': blocos,
        'ofertas': ofertas,
        'rascunho': rascunho or {},
        'rascunho_id': rascunho_id,
        'is_admin': is_admin,
        'is_superusuario': is_superusuario
    }
    
    # Organizar blocos por categoria para melhor visualização
    blocos_por_oferta = {"geral": {}}
    
    # Processar blocos gerais
    for nome_bloco, bloco in blocos.items():
        categoria = bloco.get("categoria", "geral")
        
        if categoria not in blocos_por_oferta:
            blocos_por_oferta[categoria] = {}
        
        blocos_por_oferta[categoria][nome_bloco] = bloco
        
        # Adicionar também à categoria "geral" para visualização completa
        if categoria != "geral":
            blocos_por_oferta["geral"][nome_bloco] = bloco
            
    # Processar blocos específicos de ofertas
    for oferta_nome, oferta_dados in ofertas.items():
        categoria_oferta = f"oferta_{oferta_nome}"
        
        # Garantir que a categoria existe
        if categoria_oferta not in blocos_por_oferta:
            blocos_por_oferta[categoria_oferta] = {}
        
        # Processar blocos da oferta se existirem
        if 'blocos' in oferta_dados:
            for bloco_nome, bloco_info in oferta_dados['blocos'].items():
                # Verificar se o bloco já existe
                if bloco_nome not in blocos:
                    # Verificar permissão para blocos de oferta
                    tem_permissao = is_admin or is_superusuario
                    
                    # Se não for admin/superusuário, verificar permissões específicas
                    if not tem_permissao and usuario_obj:
                        # Verificar nas permissões do usuário
                        bloco_db = BlocoProposta.query.filter_by(nome=bloco_nome).first()
                        if bloco_db and (bloco_db in usuario_obj.blocos_permitidos or bloco_db.criado_por == usuario_atual):
                            tem_permissao = True
                    
                    # Verificar se é obrigatório
                    is_obrigatorio = bloco_nome in oferta_dados.get('obrigatorios', [])
                    
                    # Adicionar se for obrigatório ou tiver permissão
                    if is_obrigatorio or tem_permissao:
                        # Criar o bloco temporariamente para exibição
                        novo_bloco = {
                            "titulo": bloco_info.get('titulo', bloco_nome.replace('_', ' ')),
                            "texto": bloco_info.get('texto', ''),
                            "obrigatorio": is_obrigatorio,
                            "categoria": categoria_oferta
                        }
                        blocos_por_oferta[categoria_oferta][bloco_nome] = novo_bloco
    
    # Adicionar a variável blocos_por_oferta aos dados do template
    template_data['blocos_por_oferta'] = blocos_por_oferta
    
    # Adicionar blocos específicos do cliente
    blocos_cliente = []
    for bloco_nome, bloco_info in blocos.items():
        if bloco_info.get('cliente_associado') == cliente:
            blocos_cliente.append(bloco_nome)
    
    template_data['blocos_cliente'] = blocos_cliente
    
    # Adicionar blocos temporários do rascunho, se existirem
    blocos_temporarios = {}
    if rascunho and 'blocos_temporarios' in rascunho:
        blocos_temporarios = rascunho['blocos_temporarios']
    
    template_data['blocos_temporarios'] = blocos_temporarios
    
    return render_template('criar_proposta.html', **template_data)

# Rota para processar a criação de uma proposta
@app.route('/criar_proposta', methods=['POST'])
@login_required
def criar_proposta():
    logger = logging.getLogger('criar_proposta')
    logger.info("="*40)
    logger.info("INÍCIO DA CRIAÇÃO DE PROPOSTA")
    logger.info("="*40)
    
    nome_cliente = request.form.get('nome_cliente', '')
    
    # Obter blocos selecionados
    blocos_selecionados = request.form.getlist('blocos_selecionados')
    
    logger.info(f"Cliente: {nome_cliente}")
    logger.info(f"Blocos selecionados: {blocos_selecionados}")
    
    # Logs detalhados para resolver problemas com o composer
    logger.info(f"Data/Hora: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info("=" * 80)
    
    try:
        # Verificar se o nome do cliente foi fornecido
        if not nome_cliente:
            flash('Por favor, informe o nome do cliente.', 'danger')
            return redirect(url_for('exibir_criar_proposta'))
        
        # Processar upload de arquivo de logo
        logo_cliente = None
        if 'logo_file' in request.files and request.files['logo_file'].filename:
            logo_file = request.files['logo_file']
            logger.info(f"Logo recebido: {logo_file.filename}")
            
            # Verificar extensão
            if logo_file and allowed_file(logo_file.filename):
                # Gerar nome único para o arquivo
                filename = secure_filename(f"{nome_cliente.replace(' ', '_')}_{uuid.uuid4()}.{logo_file.filename.rsplit('.', 1)[1].lower()}")
                logo_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                logo_file.save(logo_path)
                # Salvar o caminho relativo para facilitar a portabilidade
                logo_cliente = os.path.join('static', 'uploads', filename)
                logger.info(f"Logo salvo em: {logo_path}, caminho relativo: {logo_cliente}")
            else:
                flash('Formato de arquivo não suportado. Use PNG, JPG ou GIF.')
                logger.warning(f"Formato de arquivo não suportado: {logo_file.filename}")
                return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
        
        # Obter demais dados do formulário
        modelo_proposta = request.form.get('modelo_proposta')
        
        # Registrar blocos selecionados nos logs
        logger.info(f"Blocos selecionados: {blocos_selecionados}")
        logger.info(f"Modelo de proposta: {modelo_proposta}")
        
        # Obter oferta selecionada, se houver
        oferta_selecionada = request.form.get('oferta_selecionada')
        if oferta_selecionada and oferta_selecionada != 'nenhuma':
            logger.info(f"Oferta selecionada: {oferta_selecionada}")
        else:
            oferta_selecionada = None
            logger.info("Nenhuma oferta selecionada")
        
        # Obter a ordem dos blocos definida pelo usuário na interface
        blocos_ordem_json = request.form.get('blocos_ordem')
        blocos_ordem = None
        if blocos_ordem_json:
            try:
                blocos_ordem = json.loads(blocos_ordem_json)
                logger.info(f"Ordem personalizada dos blocos: {blocos_ordem}")
            except json.JSONDecodeError as e:
                logger.error(f"Erro ao decodificar a ordem dos blocos: {e}")
        
        # Obter dados de rascunho, se houver
        rascunho_id = request.form.get('rascunho_id')
        blocos_temporarios = None
        if rascunho_id:
            logger.info(f"Rascunho ID fornecido: {rascunho_id}")
            rascunhos = carregar_rascunhos()
            if rascunho_id in rascunhos:
                blocos_temporarios = rascunhos[rascunho_id].get('blocos_temporarios', {})
                logger.info(f"Blocos temporários encontrados: {len(blocos_temporarios) if blocos_temporarios else 0}")
        
        # Verificar se é para gerar a proposta ou salvar como rascunho
        if 'btn_salvar_rascunho' in request.form:
            logger.info("Salvando como rascunho")
            logo_atual = request.form.get('logo_atual')
            return salvar_como_rascunho(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados, rascunho_id, logo_atual, oferta_selecionada, blocos_temporarios)
        
        # Gerar proposta
        logger.info("Iniciando geração da proposta")
        output_path, filename = gerar_proposta(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados, oferta_selecionada, blocos_temporarios, blocos_ordem)
        
        # Verificar se o arquivo foi gerado
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            logger.info(f"Proposta gerada com sucesso: {output_path} (Tamanho: {file_size/1024:.2f} KB)")
        else:
            logger.error(f"Arquivo não encontrado após geração: {output_path}")
            
        # Salvar o registro da proposta no banco de dados
        try:
            # Salvar o caminho completo do arquivo, não apenas o nome
            nova_proposta = salvar_proposta(nome_cliente, output_path, session.get('usuario_logado'), blocos_selecionados, oferta_selecionada)
            logger.info(f"Proposta salva no banco de dados com ID: {nova_proposta}")
        except Exception as e:
            logger.error(f"Erro ao salvar proposta no banco de dados: {e}")
            # Continuar mesmo se falhar o registro no banco
            
        # Se houve um rascunho usado, remover após gerar a proposta
        if rascunho_id:
            try:
                remover_rascunho(rascunho_id)
                logger.info(f"Rascunho removido após geração da proposta: {rascunho_id}")
            except Exception as e:
                logger.warning(f"Não foi possível remover o rascunho: {e}")
        
        # Redirecionar para a página de visualização de proposta
        flash('Proposta gerada com sucesso!', 'success')
        return redirect(url_for('visualizar_proposta', proposta_id=nova_proposta))
    
    except Exception as e:
        logger.error(f"Erro na geração da proposta: {str(e)}", exc_info=True)
        flash(f'Erro ao gerar proposta: {str(e)}', 'danger')
        return redirect(url_for('exibir_criar_proposta'))

# Função para criar imagem placeholder
def criar_imagem_placeholder(caminho, texto, largura=500, altura=200, cor_fundo=(0, 0, 0), cor_texto=(255, 255, 255)):
    """
    Cria uma imagem placeholder com texto centralizado.
    
    Args:
        caminho: Caminho onde a imagem será salva
        texto: Texto a ser exibido na imagem
        largura: Largura da imagem em pixels
        altura: Altura da imagem em pixels
        cor_fundo: Cor de fundo da imagem (R, G, B)
        cor_texto: Cor do texto (R, G, B)
    """
    try:
        # Criar diretório se não existir
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        
        # Criar uma imagem com fundo preto
        imagem = Image.new('RGB', (largura, altura), color=cor_fundo)
        draw = ImageDraw.Draw(imagem)
        
        # Tentar carregar uma fonte, ou usar a fonte padrão
        try:
            fonte = ImageFont.truetype("arial.ttf", 20)
        except IOError:
            fonte = ImageFont.load_default()
        
        # Calcular posição do texto para centralizá-lo
        largura_texto, altura_texto = 0, 0
        try:
            # Para versões mais recentes do PIL
            left, top, right, bottom = draw.textbbox((0, 0), texto, font=fonte)
            largura_texto = right - left
            altura_texto = bottom - top
        except AttributeError:
            try:
                # Para versões intermediárias do PIL
                largura_texto, altura_texto = draw.textsize(texto, font=fonte)
            except AttributeError:
                # Fallback para valores aproximados
                largura_texto, altura_texto = len(texto) * 10, 20
        
        posicao = ((largura - largura_texto) // 2, (altura - altura_texto) // 2)
        
        # Desenhar o texto
        draw.text(posicao, texto, fill=cor_texto, font=fonte)
        
        # Salvar a imagem
        imagem.save(caminho)
        logging.info(f"Imagem placeholder criada em: {caminho}")
        return True
    except Exception as e:
        logging.error(f"Erro ao criar imagem placeholder: {e}")
        return False

def gerar_proposta(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados, oferta_selecionada=None, blocos_temporarios=None, blocos_ordem=None):
    try:
        import re
        import requests
        from docx import Document
        from docxtpl import DocxTemplate, InlineImage
        from docxcompose.composer import Composer
        from bs4 import BeautifulSoup
        
        # Definição das cores corporativas 
        SERVICE_IT_RED = RGBColor(227, 23, 54)    # Vermelho da Service IT
        SERVICE_IT_GRAY = RGBColor(83, 86, 90)    # Cinza da Service IT
        
        # Garantir que a fonte Poppins esteja disponível antes de gerar a proposta
        poppins_disponivel = garantir_fonte_poppins()
        logger.info(f"Verificação de disponibilidade da fonte Poppins concluída: {'Disponível' if poppins_disponivel else 'Indisponível'}")
        
        # Logs detalhados para composição da proposta
        logger.info(f"Gerando proposta para: {nome_cliente}")
        logger.info(f"Tipo de blocos_selecionados: {type(blocos_selecionados)}")
        logger.info(f"Total de blocos selecionados: {len(blocos_selecionados) if blocos_selecionados else 0}")
        if blocos_selecionados:
            for i, bloco in enumerate(blocos_selecionados):
                logger.info(f"  Bloco {i+1}: {bloco}")
        
        # 1. Definir o nome do arquivo de saída
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"proposta_{nome_cliente.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        logger.info(f"Arquivo de saída definido: {output_path}")
        
        # Criar diretório temporário para os arquivos intermediários
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{timestamp}")
        os.makedirs(temp_dir, exist_ok=True)
        logger.info(f"Diretório temporário criado: {temp_dir}")
        
        # 2. Usar o template da capa na raiz do projeto
        template_path = os.path.join(app.root_path, 'capa.docx')
        logging.info(f"Usando template da capa: {template_path}")
        
        # Verificar se o template existe
        if not os.path.exists(template_path):
            logger.error(f"Template da capa não encontrado em: {template_path}")
            raise FileNotFoundError(f"Template da capa não encontrado em: {template_path}")
        
        # 3. Preparar a lista de blocos a adicionar
        if blocos_selecionados:
            # Se for uma lista imutável do Flask, converter para lista Python
            if hasattr(blocos_selecionados, 'getlist'):
                logging.info("blocos_selecionados é um tipo de objeto Flask")
                blocos_a_adicionar = blocos_selecionados.getlist('blocos')
            # Se for uma lista Python regular
            elif isinstance(blocos_selecionados, (list, tuple)):
                blocos_a_adicionar = list(blocos_selecionados)
            # Se for um único valor
            else:
                blocos_a_adicionar = [blocos_selecionados]
        else:
            blocos_a_adicionar = []
        
        # Importante: Garantir que TODOS os blocos selecionados pelo usuário sejam utilizados
        # Isso é crucial para usar os blocos que o usuário escolheu na interface
        if blocos_selecionados and len(blocos_selecionados) > 0:
            blocos_a_adicionar = list(blocos_selecionados)  # Usar diretamente os blocos selecionados pelo usuário
            
        logging.info(f"Blocos selecionados (inicial): {blocos_a_adicionar}")
        
        # 4. Carregar blocos e ofertas
        blocos = carregar_blocos()
        ofertas = carregar_ofertas()
        
        # Log adicional para verificar o carregamento dos blocos
        logging.info(f"Número total de blocos carregados: {len(blocos)}")
        logging.info(f"Nomes dos blocos carregados: {list(blocos.keys())}")
        
        # 5. Adicionar blocos temporários e obrigatórios
        # Adicionar blocos temporários ao dicionário de blocos (apenas para esta proposta)
        if blocos_temporarios and isinstance(blocos_temporarios, dict):
            for nome_bloco_temp, dados_bloco_temp in blocos_temporarios.items():
                # Não sobrescrever blocos existentes
                if nome_bloco_temp not in blocos:
                    blocos[nome_bloco_temp] = {
                        'texto': dados_bloco_temp.get('texto', ''),
                        'imagem': None,
                        'obrigatorio': False,
                        'criado_por': dados_bloco_temp.get('criado_por', 'Desconhecido'),
                        'data_criacao': dados_bloco_temp.get('data_criacao', datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')),
                        'cliente_associado': nome_cliente,
                        'temporario': True  # Marcar como temporário
                    }
                    logging.info(f"Adicionado bloco temporário: {nome_bloco_temp}")
        
        # Verificar se há oferta selecionada para incluir os blocos obrigatórios
        if oferta_selecionada and oferta_selecionada in ofertas:
            blocos_obrigatorios = ofertas[oferta_selecionada].get('obrigatorios', [])
            logging.info(f"Blocos obrigatórios da oferta '{oferta_selecionada}': {blocos_obrigatorios}")
            
            # Adicionar blocos obrigatórios à lista se ainda não estiverem lá
            for bloco_obrigatorio in blocos_obrigatorios:
                if bloco_obrigatorio not in blocos_a_adicionar:
                    blocos_a_adicionar.append(bloco_obrigatorio)
                    logger.info(f"Adicionando bloco obrigatório da oferta: {bloco_obrigatorio}")
        else:
            # Se não houver oferta selecionada, usar blocos obrigatórios comuns a todas as ofertas
            blocos_obrigatorios_comuns = set()
            for oferta_nome, oferta_info in ofertas.items():
                oferta_obrigatorios = oferta_info.get('obrigatorios', [])
                if not blocos_obrigatorios_comuns:
                    # Primeira oferta - adicionamos todos os blocos
                    blocos_obrigatorios_comuns = set(oferta_obrigatorios)
                else:
                    # Próximas ofertas - mantemos apenas os blocos comuns a todas
                    blocos_obrigatorios_comuns &= set(oferta_obrigatorios)
            
            logging.info(f"Blocos obrigatórios comuns a todas as ofertas: {blocos_obrigatorios_comuns}")
            
            # Adicionar os blocos obrigatórios comuns
            for bloco_obrigatorio in blocos_obrigatorios_comuns:
                if bloco_obrigatorio not in blocos_a_adicionar:
                    blocos_a_adicionar.append(bloco_obrigatorio)
                    logging.info(f"Adicionando bloco obrigatório comum: {bloco_obrigatorio}")
        
        # Garantir que todos os blocos marcados como obrigatórios no arquivo de blocos sejam incluídos
        for bloco_nome, bloco_info in blocos.items():
            if bloco_info.get('obrigatorio', False):
                if bloco_nome not in blocos_a_adicionar:
                    blocos_a_adicionar.append(bloco_nome)
                    logging.info(f"Adicionando bloco obrigatório do arquivo blocos.json: {bloco_nome}")
        
        # 6. Garantir que todos os blocos obrigatórios existem, ou criar conteúdo padrão
        for bloco_nome in blocos_a_adicionar:
            if bloco_nome not in blocos:
                # Se o bloco não existe no arquivo de blocos, criar um conteúdo padrão
                if bloco_nome == "Termo_de_Confidencialidade":
                    texto_padrao = "<p>Este documento contém informações confidenciais. A divulgação, distribuição ou reprodução deste documento sem autorização prévia por escrito é estritamente proibida.</p>"
                elif bloco_nome == "Folha_de_Rosto":
                    texto_padrao = f"<p>Proposta de Serviços para {nome_cliente}</p><p>Data de criação: {datetime.datetime.now().strftime('%d/%m/%Y')}</p>"
                elif bloco_nome == "Controle_de_Versao":
                    texto_padrao = f"<p>Versão 1.0 - {datetime.datetime.now().strftime('%d/%m/%Y')} - Documento inicial</p>"
                elif bloco_nome == "Sumario_Executivo":
                    texto_padrao = f"<p>A Service IT tem o prazer de apresentar esta proposta para {nome_cliente}, desenvolvida para atender às necessidades específicas da sua organização.</p>"
                elif bloco_nome == "Sobre_a_Service_IT":
                    texto_padrao = "<p>A Service IT é uma empresa líder em soluções de tecnologia, oferecendo serviços de alta qualidade e confiabilidade para diversos setores do mercado.</p>"
                elif bloco_nome == "Visao_Geral_dos_Servicos":
                    texto_padrao = "<p>Nesta seção, apresentamos uma visão geral dos serviços oferecidos pela Service IT, destacando os principais benefícios e diferenciais.</p>"
                elif bloco_nome == "Solucao_Proposta":
                    texto_padrao = f"<p>A solução proposta para {nome_cliente} foi desenvolvida considerando as necessidades específicas da sua organização e os desafios atuais do seu negócio.</p>"
                elif bloco_nome == "Exclusoes":
                    texto_padrao = "<p>Esta proposta não inclui serviços ou produtos que não estejam explicitamente mencionados neste documento.</p>"
                elif bloco_nome == "Transicao":
                    texto_padrao = "<p>A Service IT implementará um plano de transição estruturado para garantir a implementação suave e eficiente dos serviços propostos.</p>"
                elif bloco_nome == "Termo_de_Aceite_da_Proposta":
                    texto_padrao = f"<p>Ao assinar este documento, {nome_cliente} concorda com os termos e condições desta proposta.</p>"
                else:
                    texto_padrao = f"<p>Conteúdo para o bloco '{bloco_nome.replace('_', ' ').title()}' será definido.</p>"
                
                blocos[bloco_nome] = {
                    'texto': texto_padrao,
                    'imagem': None,
                    'obrigatorio': True,
                    'criado_por': 'Sistema',
                    'data_criacao': datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                    'cliente_associado': '',
                    'temporario': False
                }
                logging.info(f"Criado bloco padrão para o obrigatório: {bloco_nome}")
        
        # 7. Adicionar blocos específicos do cliente se não estiverem já incluídos
        for bloco_nome, bloco_info in blocos.items():
            if bloco_info.get('cliente_associado') == nome_cliente and bloco_nome not in blocos_a_adicionar:
                blocos_a_adicionar.append(bloco_nome)
                logging.info(f"Adicionando bloco específico do cliente: {bloco_nome}")
        
        # 8. Ordenar os blocos
        ordem_padrao = [
            "Termo_de_Confidencialidade",
            "Folha_de_Rosto",
            "Controle_de_Versao",
            "Sumario_Executivo",
            "Sobre_a_Service_IT",
            "Visao_Geral_dos_Servicos",
            "Solucao_Proposta",
            "Exclusoes",
            "Transicao",
            "Termo_de_Aceite_da_Proposta"
        ]
        
        # MODIFICAÇÃO: Usar a ordem de blocos definida pelo usuário na interface, se disponível
        if blocos_ordem and isinstance(blocos_ordem, list) and len(blocos_ordem) > 0:
            logging.info(f"Usando ordem de blocos personalizada definida pelo usuário: {blocos_ordem}")
            
            # Criar uma lista ordenada com os blocos na ordem definida pelo usuário
            blocos_a_adicionar_ordenados = []
            
            # Primeiro adicionar os blocos na ordem definida pelo usuário
            for bloco_nome in blocos_ordem:
                if bloco_nome in blocos_a_adicionar:
                    blocos_a_adicionar_ordenados.append(bloco_nome)
                    logging.info(f"Adicionando bloco da hierarquia personalizada: {bloco_nome}")
            
            # Adicionar outros blocos selecionados que não estão na hierarquia personalizada
            for bloco_nome in blocos_a_adicionar:
                if bloco_nome not in blocos_a_adicionar_ordenados:
                    blocos_a_adicionar_ordenados.append(bloco_nome)
                    logging.info(f"Adicionando bloco selecionado não presente na hierarquia: {bloco_nome}")
            
            # Atualizar a lista de blocos a adicionar
            blocos_a_adicionar = blocos_a_adicionar_ordenados
            
            logging.info(f"Ordem final dos blocos (com hierarquia personalizada): {blocos_a_adicionar}")
        else:
            # Fallback para o comportamento anterior se não houver hierarquia personalizada
            if oferta_selecionada and oferta_selecionada in ofertas and 'obrigatorios' in ofertas[oferta_selecionada]:
                ordem_padrao = ofertas[oferta_selecionada]['obrigatorios']
                logging.info(f"Usando ordem de blocos da oferta {oferta_selecionada}")
            
            # Separar blocos obrigatórios e outros blocos
            blocos_ordenados = []
            blocos_nao_obrigatorios = []
            
            # MODIFICAÇÃO: Primeiro preservar todos os blocos selecionados pelo usuário
            # Copiar a lista atual de blocos para garantir que nenhum seja perdido
            blocos_a_adicionar_original = blocos_a_adicionar.copy()
            
            # Primeiro, adicionar os blocos na ordem padrão
            for bloco_nome in ordem_padrao:
                if bloco_nome in blocos_a_adicionar:
                    blocos_ordenados.append(bloco_nome)
            
            # Depois, adicionar os blocos restantes
            for bloco_nome in blocos_a_adicionar:
                if bloco_nome not in blocos_ordenados:
                    blocos_nao_obrigatorios.append(bloco_nome)
            
            # Ordenar os blocos não obrigatórios alfabeticamente
            blocos_nao_obrigatorios.sort()
            
            # Combinar as listas
            blocos_a_adicionar = blocos_ordenados + blocos_nao_obrigatorios
            
            # MODIFICAÇÃO: Verificar e garantir que todos os blocos originais estejam ainda presentes
            # Adicionar de volta quaisquer blocos que possam ter sido perdidos no processo de ordenação
            for bloco_nome in blocos_a_adicionar_original:
                if bloco_nome not in blocos_a_adicionar:
                    blocos_a_adicionar.append(bloco_nome)
                    logging.info(f"Recuperando bloco que seria perdido: {bloco_nome}")
            
            logging.info(f"Ordem final dos blocos (usando ordem padrão): {blocos_a_adicionar}")
        
        # 9. MÉTODO CORRIGIDO: Criar documento principal com a capa
        doc_principal = Document()
        
        # Definir estilos para o documento principal
        styles = doc_principal.styles
        
        # Estilo de parágrafo normal
        if 'Normal' in styles:
            style_normal = styles['Normal']
            font = style_normal.font
            font.name = get_font_name('Poppins')
            font.size = Pt(11)
            
        # Estilo de cabeçalhos
        if 'Heading 2' in styles:
            style_heading2 = styles['Heading 2']
            font = style_heading2.font
            font.name = get_font_name('Poppins')
            font.size = Pt(14)
            font.bold = True
            font.color.rgb = SERVICE_IT_RED
        
        # 10. Gerar a CAPA com DocxTemplate e incorporar ao documento
        capa_doc = DocxTemplate(template_path)
        
        # Preparar o contexto para substituição de variáveis
        context = {
            'NOME_CLIENTE': nome_cliente,
            'nome_cliente': nome_cliente,
            'Nome_Cliente': nome_cliente,
        }
        
        # Adicionar logo ao contexto
        logo_path = None
        if logo_cliente:
            logging.info(f"Verificando logo do cliente: {logo_cliente}")
            logo_path = logo_cliente
            if not os.path.isabs(logo_cliente):
                logo_path = os.path.join(app.root_path, logo_cliente)
            
            logging.info(f"Caminho completo do logo: {logo_path}")
        
        if logo_path and os.path.exists(logo_path):
            # Carregar a imagem para verificar dimensões
            try:
                from PIL import Image
                img = Image.open(logo_path)
                largura, altura = img.size
                proporcao = largura / altura
                
                # Ajustar dimensões com base na proporção da imagem
                if proporcao > 2:  # Logo muito larga e baixa
                    context['logo_cliente'] = InlineImage(capa_doc, logo_path, width=Inches(2.2))
                elif proporcao < 0.8:  # Logo muito alta e estreita
                    context['logo_cliente'] = InlineImage(capa_doc, logo_path, width=Inches(1.8), height=Inches(1.8))
                else:  # Logo com proporção normal
                    context['logo_cliente'] = InlineImage(capa_doc, logo_path, width=Inches(2.2))
                
                logging.info(f"Logo do cliente adicionado ao contexto: {logo_path} (proporção: {proporcao:.2f})")
            except Exception as e:
                # Fallback em caso de erro na análise da imagem
                context['logo_cliente'] = InlineImage(capa_doc, logo_path, width=Inches(2.0), height=Inches(1.5))
                logging.warning(f"Erro ao analisar dimensões da logo, usando tamanho padrão: {str(e)}")
        else:
            logging.warning("Nenhum logo do cliente fornecido")
            logo_padrao = os.path.join(app.root_path, 'static', 'img', 'placeholder-logo.png')
            if os.path.exists(logo_padrao):
                context['logo_cliente'] = InlineImage(capa_doc, logo_padrao, width=Inches(2.2))
                logging.info("Logo padrão adicionado ao contexto")
        
        # Renderizar a capa e salvar
        capa_doc.render(context)
        capa_path = os.path.join(temp_dir, f"capa_{timestamp}.docx")
        capa_doc.save(capa_path)
        logger.info(f"Capa gerada e salva em: {capa_path}")
        
        # Inicializar o documento final com a capa
        documento_final_path = os.path.join(temp_dir, f"final_{timestamp}.docx")
        documento_final = Document(capa_path)
        
        # Adicionar uma quebra de página após a capa
        documento_final.add_page_break()
        logger.info("Adicionada quebra de página após a capa")
        
        # Função para substituir variáveis em texto HTML
        def substituir_variaveis_html(texto_html, nome_cliente):
            formatos_placeholder = ["{{NOME_CLIENTE}}", "{{nome_cliente}}"]
            formatos_placeholder_compat = ["[[NOME_CLIENTE]]", "NOME_CLIENTE", "[[nome_cliente]]", "nome_cliente"]
            
            for formato in formatos_placeholder + formatos_placeholder_compat:
                if formato in texto_html:
                    texto_html = texto_html.replace(formato, nome_cliente)
            return texto_html
        
        # 11. Adicionar cada bloco diretamente ao documento final
        arquivos_temporarios = [capa_path]  # Lista para controlar arquivos temporários
        
        # Logs temporários para debug
        logger.info(f"Blocos que serão adicionados: {blocos_a_adicionar}")
        logger.info(f"Total de blocos disponíveis: {len(blocos)}")
        
        # Adicionar espaço extra antes do primeiro bloco
        espaco_inicial = documento_final.add_paragraph()
        espaco_inicial.add_run("\n\n")  # espaço adicional antes do primeiro bloco
        logger.info("Adicionado espaço extra antes do primeiro bloco")
        
        for indice, bloco_nome in enumerate(blocos_a_adicionar, 1):
            try:
                logger.info(f"Processando bloco {indice}: {bloco_nome}")
                
                # MODIFICADO: Não adicionar quebra de página automática entre blocos
                # Apenas adicionar espaço entre os blocos com um parágrafo
                if indice > 1:
                    # Adicionar espaço entre blocos
                    espaco_entre = documento_final.add_paragraph()
                    espaco_entre.add_run("\n\n")  # espaço adicional entre blocos
                    espaco_entre.space_after = Pt(24)  # adicionar espaço após o parágrafo
                    logger.info(f"Adicionado espaço entre os blocos {indice-1} e {indice}")
                
                # Formatar título do bloco
                nome_bloco_formatado = bloco_nome.replace('_', ' ').title()
                nome_bloco_formatado = re.sub(r'^\s*\d+(\.\d+)*\s+', '', nome_bloco_formatado).strip()
                
                # Adicionar título como parágrafo estilizado
                p = documento_final.add_paragraph()
                # p.style = 'Heading 2'  # Removido para evitar numeração automática
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.space_before = Pt(18)
                p.space_after = Pt(12)
                
                run = p.add_run(f"{indice}. {nome_bloco_formatado}")
                run.bold = True
                run.font.name = get_font_name('Poppins')
                run.font.size = Pt(14)
                run.font.color.rgb = SERVICE_IT_RED
                
                # Verificar se o bloco existe na biblioteca
                if bloco_nome in blocos:
                    bloco = blocos[bloco_nome]
                    
                    # Processar texto do bloco
                    texto_bloco = bloco.get('texto', '')
                    
                    # Substituir placeholders no texto
                    if texto_bloco:
                        # Remover prefixos de numeração
                        texto_bloco = re.sub(r'<p>\s*\d+(?:\.\d+)*\.?\s*', '<p>', texto_bloco)
                        
                        # Substituir variáveis do cliente
                        texto_bloco = substituir_variaveis_html(texto_bloco, nome_cliente)
                        
                        # Usar BeautifulSoup para processar HTML
                        soup = BeautifulSoup(texto_bloco, 'html.parser')
                        
                        if soup and len(soup.contents) > 0:
                            # Processar cada elemento HTML
                            for element in soup.children:
                                try:
                                    # Processar texto simples
                                    if element.name is None:
                                        if element.strip():
                                            p = documento_final.add_paragraph()
                                            run = p.add_run(element.strip())
                                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                            run.font.name = get_font_name('Poppins')
                                            run.font.size = Pt(11)
                                            run.font.color.rgb = SERVICE_IT_GRAY
                                    
                                    # Processar parágrafos
                                    elif element.name == 'p':
                                        p = documento_final.add_paragraph()
                                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                        
                                        # Processar o conteúdo do parágrafo
                                        for child in element.children:
                                            # Texto simples
                                            if child.name is None:
                                                run = p.add_run(child.string)
                                                run.font.name = get_font_name('Poppins')
                                                run.font.size = Pt(11)
                                                run.font.color.rgb = SERVICE_IT_GRAY
                                            
                                            # Negrito
                                            elif child.name == 'strong' or child.name == 'b':
                                                run = p.add_run(child.get_text())
                                                run.bold = True
                                                run.font.name = get_font_name('Poppins')
                                                run.font.size = Pt(11)
                                                run.font.color.rgb = SERVICE_IT_RED
                                            
                                            # Itálico
                                            elif child.name == 'em' or child.name == 'i':
                                                run = p.add_run(child.get_text())
                                                run.italic = True
                                                run.font.name = get_font_name('Poppins')
                                                run.font.size = Pt(11)
                                                run.font.color.rgb = SERVICE_IT_GRAY
                                            
                                            # Imagens
                                            elif child.name == 'img':
                                                try:
                                                    img_src = child.get('src', '')
                                                    if img_src:
                                                        # Processamento de caminhos de imagem
                                                        img_path = None
                                                        
                                                        # Tratar imagens em formato base64
                                                        if img_src.startswith("data:image"):
                                                            # É um base64 embutido
                                                            import base64
                                                            import uuid
                                                            header, base64_data = img_src.split(',', 1)
                                                            image_data = base64.b64decode(base64_data)
                                                            img_temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"img_{uuid.uuid4()}.png")
                                                            with open(img_temp_path, "wb") as f:
                                                                f.write(image_data)
                                                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                            run = p.add_run()
                                                            run.add_picture(img_temp_path, width=Inches(5))
                                                            logger.info(f"Imagem base64 adicionada ao documento: {img_temp_path}")
                                                            # Adicionar à lista de arquivos temporários para limpar depois
                                                            arquivos_temporarios.append(img_temp_path)
                                                            continue
                                                        
                                                        # Converter URLs relativas para caminhos absolutos
                                                        if img_src.startswith('/static/'):
                                                            # URL relativa ao root
                                                            img_path = os.path.join(app.root_path, img_src.lstrip('/'))
                                                        elif img_src.startswith('static/'):
                                                            # URL relativa
                                                            img_path = os.path.join(app.root_path, img_src)
                                                        elif 'uploads/' in img_src:
                                                            # Imagem em uploads
                                                            filename = os.path.basename(img_src)
                                                            img_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                                                        else:
                                                            # Tentar outros caminhos possíveis
                                                            potential_paths = [
                                                                img_src,
                                                                os.path.join(app.root_path, img_src),
                                                                os.path.join(app.root_path, 'static', 'uploads', os.path.basename(img_src)),
                                                                os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(img_src))
                                                            ]
                                                            
                                                            for path in potential_paths:
                                                                if os.path.exists(path):
                                                                    img_path = path
                                                                    break
                                                        
                                                        # Se encontrou um caminho válido, adicionar a imagem
                                                        if img_path and os.path.exists(img_path):
                                                            # Centralizar o parágrafo
                                                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                            
                                                            # Adicionar a imagem com tamanho adequado
                                                            run = p.add_run()
                                                            # Usar um tamanho razoável para a imagem (ajuste conforme necessário)
                                                            run.add_picture(img_path, width=Inches(5))
                                                            logger.info(f"Imagem adicionada ao documento: {img_path}")
                                                        else:
                                                            logger.error(f"Imagem não encontrada: {img_src}")
                                                except Exception as e:
                                                    logger.error(f"Erro ao processar imagem: {str(e)}")
                                            
                                            # Outros elementos
                                            else:
                                                run = p.add_run(child.get_text())
                                                run.font.name = get_font_name('Poppins')
                                                run.font.size = Pt(11)
                                                run.font.color.rgb = SERVICE_IT_GRAY
                                    
                                except Exception as e:
                                    logging.error(f"Erro ao processar elemento HTML: {str(e)}")
                                    continue
                        else:
                            # Adicionar um parágrafo genérico se o bloco não tiver texto
                            p = documento_final.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            run = p.add_run(f"Conteúdo a ser definido para '{bloco_nome.replace('_', ' ').title()}'")
                            run.italic = True
                            run.font.name = get_font_name('Poppins')
                            run.font.size = Pt(11)
                            run.font.color.rgb = SERVICE_IT_GRAY
                    else:
                        # Adicionar um parágrafo genérico se o bloco não tiver texto
                        p = documento_final.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run = p.add_run(f"Conteúdo a ser definido para '{bloco_nome.replace('_', ' ').title()}'")
                        run.italic = True
                        run.font.name = get_font_name('Poppins')
                        run.font.size = Pt(11)
                        run.font.color.rgb = SERVICE_IT_GRAY
                else:
                    # Adicionar um parágrafo genérico se o bloco não existir na biblioteca
                    logging.warning(f"Bloco não encontrado: {bloco_nome}")
                    p = documento_final.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(f"[Bloco '{bloco_nome.replace('_', ' ').title()}' não encontrado na biblioteca]")
                    run.italic = True
                    run.font.name = get_font_name('Poppins')
                    run.font.size = Pt(11)
                    run.font.color.rgb = SERVICE_IT_GRAY

                # Adicionar espaço entre blocos
                espaco_entre_blocos = documento_final.add_paragraph()
                espaco_entre_blocos.add_run("\n")  # uma linha em branco
                
            except Exception as e:
                logger.error(f"Erro ao processar bloco {bloco_nome}: {str(e)}", exc_info=True)
                # Continuar para o próximo bloco
        
        # 12. Salvar o documento final
        documento_final.save(output_path)
        logger.info(f"Documento final salvo em: {output_path}")
        
        # 13. Limpar arquivos temporários
        for arquivo in arquivos_temporarios:
            try:
                os.remove(arquivo)
                logger.info(f"Arquivo temporário removido: {arquivo}")
            except Exception as e:
                logger.warning(f"Erro ao remover arquivo temporário {arquivo}: {str(e)}")
        
        # Remover diretório temporário
        try:
            os.rmdir(temp_dir)
            logger.info(f"Diretório temporário removido: {temp_dir}")
        except Exception as e:
            logger.warning(f"Erro ao remover diretório temporário {temp_dir}: {str(e)}")
        
        # Retornar o caminho e o nome do arquivo
        return output_path, filename
        
    except Exception as e:
        logger.error(f"Erro ao gerar proposta: {str(e)}", exc_info=True)
        # Limpar eventuais arquivos temporários se ocorrer erro
        try:
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir)
                logger.info(f"Diretório temporário removido após erro: {temp_dir}")
        except Exception as cleanup_error:
            logger.warning(f"Erro ao limpar diretório temporário: {str(cleanup_error)}")
        raise

# Função auxiliar para substituir variáveis nos parágrafos
def substituir_variaveis(paragraph, placeholder, valor):
    if not paragraph or not placeholder or not valor:
        return
    
    # Padronizando para o formato {{}}
    # Verificar se o placeholder já está no formato {{}}
    if placeholder.startswith('{{') and placeholder.endswith('}}'):
        formato_padrao = placeholder
    else:
        # Converter para o formato {{}}
        formato_padrao = '{{' + placeholder.replace('[[', '').replace(']]', '').replace('{{', '').replace('}}', '') + '}}'
    
    # Lista de possíveis formatos de placeholder
    formatos_placeholder = [
        formato_padrao,  # Formato padrão (ex: {{NOME_CLIENTE}})
        formato_padrao.lower(),  # Minúsculas (ex: {{nome_cliente}})
    ]
    
    # Formatos antigos para compatibilidade
    formatos_compat = [
        placeholder,  # Formato original (ex: [[NOME_CLIENTE]])
        placeholder.replace('[[', '').replace(']]', ''),  # Sem delimitadores (ex: NOME_CLIENTE)
        placeholder.lower(),  # Minúsculas (ex: [[nome_cliente]])
    ]
    
    # Substituir cada formato possível
    for formato in formatos_placeholder + formatos_compat:
        if formato in paragraph.text:
            for run in paragraph.runs:
                if formato in run.text:
                    run.text = run.text.replace(formato, valor)

def substituir_variaveis_robusto(doc, placeholder, valor):
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder in full_text:
            # Limpar o texto existente
            for run in paragraph.runs:
                run.text = ''
            # Recriar o texto com o valor substituído
            if placeholder == '{{logo_cliente}}' and os.path.exists(valor):
                paragraph.add_run().add_picture(valor, width=Inches(2))
            else:
                paragraph.add_run(full_text.replace(placeholder, valor))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_variaveis_robusto(cell, placeholder, valor)

# Rota padrão - redireciona para login se não estiver autenticado
@app.route('/')
@login_required
def index():
    return redirect(url_for('dashboard'))

# Rota para a página inicial após o login
@app.route('/home')
@login_required
def home():
    return redirect(url_for('index'))

# Rota para gerenciar usuários
@app.route('/gerenciar_usuarios', methods=['GET', 'POST'])
@admin_required
def gerenciar_usuarios():
    """
    Função para gerenciar usuários (adicionar, remover, ativar, desativar)
    """
    # Carregar usuários existentes
    usuarios_carregados = {}
    perfis_disponiveis = ["AM (Account Manager)", "Comercial Paraná", "Comercial Rio de Janeiro", 
                         "Comercial Rio Grande do Sul", "Comercial São Paulo", "Sales Engineer (Arquitetos)",
                         "Marketing", "RH", "Temporário"]
    
    try:
        # Tentar carregar usuários do banco de dados
        usuarios_db = Usuario.query.all()
        for usuario in usuarios_db:
            tipo_usuario = "usuario"
            
            # Obter o perfil do usuário
            perfil = Perfil.query.get(usuario.id_perfil)
            if perfil:
                if perfil.nome == "Governança":
                    tipo_usuario = "admin"
                else:
                    tipo_usuario = perfil.nome
                    
            status_temp = ""
            if perfil and perfil.acesso_temporario and usuario.data_expiracao:
                if usuario.is_acesso_expirado():
                    status_temp = " (Expirado)"
                else:
                    dias_restantes = (usuario.data_expiracao - datetime.datetime.utcnow()).days
                    status_temp = f" (Expira em {dias_restantes} dias)"
                    
            # Determinar status em texto
            status_texto = "Ativo" if usuario.status == 1 else "Inativo"
            if status_temp:
                status_texto += status_temp
                    
            usuarios_carregados[usuario.login] = {
                "nome": usuario.nome,
                "senha": usuario.senha,
                "tipo": tipo_usuario,
                "perfil": perfil.nome if perfil else "",
                "status": usuario.status,
                "status_texto": status_texto,
                "superusuario": usuario.superusuario,
                "blocos_permitidos": [] # Será preenchido posteriormente
            }
            
            # Carregar blocos permitidos para este usuário
            if not usuario.superusuario and tipo_usuario != 'admin':
                try:
                    blocos_permitidos = UsuarioBloco.query.filter_by(id_usuario=usuario.id).all()
                    usuarios_carregados[usuario.login]["blocos_permitidos"] = [
                        bloco.nome_bloco for bloco in blocos_permitidos
                    ]
                except Exception as e:
                    logging.error(f"Erro ao carregar blocos permitidos para usuário {usuario.login}: {e}")
    except Exception as e:
        logging.error(f"Erro ao carregar usuários do banco: {e}")
        # Se houver erro no banco, carrega do JSON como fallback
        usuarios_carregados = carregar_usuarios()
    
    # Carregar blocos para seleção de permissões
    blocos_db = {}
    try:
        # Carregar do banco de dados
        for bloco in BlocoProposta.query.all():
            blocos_db[bloco.nome] = {
                "titulo": bloco.titulo or bloco.nome.replace('_', ' '),
                "texto": bloco.texto or "",
                "obrigatorio": bloco.obrigatorio,
                "criado_por": bloco.criado_por,
                "data_criacao": bloco.data_criacao.strftime('%d/%m/%Y %H:%M:%S') if bloco.data_criacao else "",
                "categoria": bloco.categoria or "geral"  # Usar categoria do bloco ou padrão "geral"
            }
    except Exception as e:
        logging.error(f"Erro ao carregar blocos do banco: {e}")
        # Fallback para JSON em caso de erro
        blocos_db = carregar_blocos()
        
        # Garantir que todos os blocos tenham título
        for nome_bloco, dados_bloco in blocos_db.items():
            if 'titulo' not in dados_bloco:
                blocos_db[nome_bloco]['titulo'] = nome_bloco.replace('_', ' ')
                logging.info(f"Adicionado título para bloco {nome_bloco}")
            
            # Garantir que exista o campo usuários_permitidos se for carregado do JSON
            if 'usuarios_permitidos' not in dados_bloco:
                blocos_db[nome_bloco]['usuarios_permitidos'] = []
        
        logging.info(f"Fallback para JSON: {len(blocos_db)} blocos carregados com sucesso")
        logging.info(f"Blocos carregados: {list(blocos_db.keys())}")
    
    # Debug para verificar blocos
    logging.info(f"Total de blocos carregados: {len(blocos_db)}")
    for nome, dados in blocos_db.items():
        logging.info(f"Bloco encontrado: {nome} - {dados.get('titulo', '')}")
    
    # Se não houver blocos na base ou no JSON, tentar usar diretamente o conteúdo do arquivo JSON
    if not blocos_db:
        try:
            # Verificar se o arquivo existe
            blocos_json_path = os.path.join('data', 'blocos.json')
            if os.path.exists(blocos_json_path):
                with open(blocos_json_path, 'r', encoding='utf-8') as f:
                    blocos_db = json.load(f)
                    logging.info(f"Blocos carregados diretamente do arquivo JSON: {len(blocos_db)}")
        except Exception as e:
            logging.error(f"Erro ao carregar blocos diretamente do JSON: {e}")
    
    # Se não houver blocos no banco ou no JSON, criar alguns blocos padrão
    if not blocos_db:
        logging.warning("Nenhum bloco encontrado! Criando blocos padrão temporários para exibição")
        blocos_db = {
            "Bloco_Padrao_1": {
                "titulo": "Bloco Padrão 1",
                "texto": "Texto padrão 1",
                "obrigatorio": True,
                "criado_por": "sistema",
                "data_criacao": datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            },
            "Bloco_Padrao_2": {
                "titulo": "Bloco Padrão 2",
                "texto": "Texto padrão 2",
                "obrigatorio": False,
                "criado_por": "sistema",
                "data_criacao": datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            }
        }
    
    if request.method == 'POST':
        acao = request.form.get('acao')
        
        if acao == 'adicionar':
            # Adicionar novo usuário
            novo_usuario = request.form.get('novo_usuario')
            nova_senha = request.form.get('nova_senha')
            tipo_usuario = request.form.get('tipo_usuario')
            nome_usuario = request.form.get('nome_usuario', novo_usuario)
            is_superusuario = request.form.get('is_superusuario') == 'on'
            blocos_permitidos = request.form.getlist('blocos_permitidos')
            
            # Verificar se é temporário e adicionar data de expiração
            is_temporario = tipo_usuario == 'temporario'
            data_expiracao = None
            if is_temporario:
                # Adicionar 24 horas à data atual para usuários temporários
                data_expiracao = datetime.datetime.utcnow() + timedelta(hours=24)
            
            if not novo_usuario or not nova_senha or not tipo_usuario:
                flash('Todos os campos são obrigatórios.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            tipos_validos = ['admin', 'am', 'marketing', 'rh', 'temporario', 'comercialpr', 'comercialrj', 'comercialrs', 'comercialsp', 'se']
            
            if tipo_usuario not in tipos_validos:
                flash('Tipo de usuário inválido.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            # Se for admin, verificar se a senha de admin foi fornecida
            if tipo_usuario == 'admin':
                senha_admin = request.form.get('senha_admin', '')
                if not senha_admin:
                    flash('Para criar um administrador, é necessário fornecer a senha de administrador.', 'danger')
                    return redirect(url_for('gerenciar_usuarios'))
                
                # Verificar se a senha de admin está correta (usando o usuário atual)
                usuario_admin = Usuario.query.filter_by(login=session['usuario_logado']).first()
                if not usuario_admin or not check_password_hash(usuario_admin.senha, senha_admin):
                    flash('Senha de administrador incorreta.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            try:
                # Verificar se o usuário já existe
                usuario_existente = Usuario.query.filter_by(login=novo_usuario).first()
                if usuario_existente:
                    flash('Este nome de usuário já está em uso.', 'danger')
                    return redirect(url_for('gerenciar_usuarios'))
                    
                # Obter o perfil correto com base no tipo_usuario
                perfil = None
                if tipo_usuario == 'admin':
                    perfil = Perfil.query.filter_by(nome="Governança").first()
                elif tipo_usuario == 'marketing':
                    perfil = Perfil.query.filter_by(nome="Marketing").first()
                    if not perfil:
                        perfil = Perfil(nome="Marketing", descricao="Equipe de Marketing")
                        db.session.add(perfil)
                        db.session.commit()
                elif tipo_usuario == 'rh':
                    perfil = Perfil.query.filter_by(nome="RH").first()
                    if not perfil:
                        perfil = Perfil(nome="RH", descricao="Recursos Humanos")
                        db.session.add(perfil)
                        db.session.commit()
                elif tipo_usuario == 'temporario':
                    perfil = Perfil.query.filter_by(nome="Temporário").first()
                    if not perfil:
                        perfil = Perfil(nome="Temporário", descricao="Acesso temporário por 24h", acesso_temporario=True)
                        db.session.add(perfil)
                        db.session.commit()
                else:
                    # Para outros tipos (se, am, comercial*)
                    perfil = Perfil.query.filter_by(nome=tipo_usuario.upper()).first()
                    if not perfil:
                        perfil = Perfil(nome=tipo_usuario.upper())
                        db.session.add(perfil)
                        db.session.commit()
                    
                # Criar o novo usuário no banco de dados
                senha_hash = generate_password_hash(nova_senha)
                
                novo_usuario_db = Usuario(
                    login=novo_usuario,
                    senha=senha_hash,
                    nome=nome_usuario,
                    status=1,
                    id_perfil=perfil.id,
                    superusuario=is_superusuario,
                    data_criacao=datetime.datetime.utcnow(),
                    data_expiracao=data_expiracao
                )
                
                # Adicionar permissões de blocos
                if blocos_permitidos:
                    for nome_bloco in blocos_permitidos:
                        bloco = BlocoProposta.query.filter_by(nome=nome_bloco).first()
                        if bloco:
                            novo_usuario_db.blocos_permitidos.append(bloco)
                
                db.session.add(novo_usuario_db)
                db.session.commit()
                
                # Backup em JSON
                usuarios = carregar_usuarios()
                usuarios[novo_usuario] = {
                    "senha": nova_senha,  # Nota: no JSON fica em texto plano (legado)
                    "tipo": tipo_usuario,
                    "nome": nome_usuario,
                    "status": 1,
                    "superusuario": is_superusuario,
                    "data_criacao": datetime.datetime.utcnow().strftime('%d/%m/%Y %H:%M:%S'),
                    "data_expiracao": data_expiracao.strftime('%d/%m/%Y %H:%M:%S') if data_expiracao else None
                }
                salvar_usuarios(usuarios)
                
                flash('Usuário adicionado com sucesso!', 'success')
            except Exception as e:
                flash(f'Erro ao adicionar usuário: {str(e)}', 'danger')
                logging.error(f"Erro ao adicionar usuário: {e}")
                
        elif acao == 'remover':
            # Remover usuário
            usuario_remover = request.form.get('usuario_remover')
            
            if usuario_remover == 'admin':
                flash('Não é possível remover o usuário admin.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            try:
                usuario = Usuario.query.filter_by(login=usuario_remover).first()
                if usuario:
                    db.session.delete(usuario)
                    db.session.commit()
                    
                    # Backup em JSON
                    usuarios = carregar_usuarios()
                    if usuario_remover in usuarios:
                        del usuarios[usuario_remover]
                    salvar_usuarios(usuarios)
                
                flash('Usuário removido com sucesso!', 'success')
            except Exception as e:
                flash(f'Erro ao remover usuário: {str(e)}', 'danger')
                logging.error(f"Erro ao remover usuário: {e}")
        
        elif acao == 'alternar_status':
            # Alternar status do usuário (ativar/desativar)
            usuario_alternar = request.form.get('usuario_alternar')
            
            if usuario_alternar == 'admin':
                flash('Não é possível desativar o usuário admin.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            try:
                usuario = Usuario.query.filter_by(login=usuario_alternar).first()
                if usuario:
                    # Alternar status (0 para 1, 1 para 0)
                    usuario.status = 1 if usuario.status == 0 else 0
                    db.session.commit()
                    
                    # Backup em JSON
                    usuarios = carregar_usuarios()
                    if usuario_alternar in usuarios:
                        usuarios[usuario_alternar]['status'] = 1 if usuarios[usuario_alternar].get('status', 0) == 0 else 0
                        salvar_usuarios(usuarios)
                
                status_texto = "ativado" if usuario.status == 1 else "desativado"
                flash(f'Usuário {status_texto} com sucesso!', 'success')
            except Exception as e:
                flash(f'Erro ao alternar status do usuário: {str(e)}', 'danger')
                logging.error(f"Erro ao alternar status do usuário: {e}")
        
        elif acao == 'alterar_tipo':
            # Alterar tipo do usuário
            usuario_alterar = request.form.get('usuario_alterar')
            novo_tipo = request.form.get('novo_tipo')
            is_superusuario = request.form.get('is_superusuario') == 'on'
            blocos_permitidos = request.form.getlist('blocos_permitidos')
            
            if usuario_alterar == 'admin' and novo_tipo != 'admin':
                flash('Não é possível alterar o tipo do usuário admin.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            tipos_validos = ['admin', 'am', 'marketing', 'rh', 'temporario', 'comercialpr', 'comercialrj', 'comercialrs', 'comercialsp', 'se']
            
            if novo_tipo not in tipos_validos:
                flash('Tipo de usuário inválido.', 'danger')
                return redirect(url_for('gerenciar_usuarios'))
            
            # Se estiver alterando para admin, verificar a senha
            if novo_tipo == 'admin':
                senha_admin = request.form.get('senha_admin', '')
                if not senha_admin:
                    flash('Para criar um administrador, é necessário fornecer a senha de administrador.', 'danger')
                    return redirect(url_for('gerenciar_usuarios'))
                
                # Verificar se a senha de admin está correta (usando o usuário atual)
                usuario_admin = Usuario.query.filter_by(login=session['usuario_logado']).first()
                if not usuario_admin or not check_password_hash(usuario_admin.senha, senha_admin):
                    flash('Senha de administrador incorreta.', 'danger')
                    return redirect(url_for('gerenciar_usuarios'))
            
            try:
                usuario = Usuario.query.filter_by(login=usuario_alterar).first()
                if usuario:
                    # Verificar se o usuário é temporário
                    is_temporario = novo_tipo == 'temporario'
                    data_expiracao = None
                    if is_temporario:
                        # Adicionar 24 horas à data atual para usuários temporários
                        data_expiracao = datetime.datetime.utcnow() + timedelta(hours=24)
                    
                    # Obter o perfil correto com base no novo_tipo
                    perfil = None
                    if novo_tipo == 'admin':
                        perfil = Perfil.query.filter_by(nome="Governança").first()
                    elif novo_tipo == 'marketing':
                        perfil = Perfil.query.filter_by(nome="Marketing").first()
                        if not perfil:
                            perfil = Perfil(nome="Marketing", descricao="Equipe de Marketing")
                            db.session.add(perfil)
                            db.session.commit()
                    elif novo_tipo == 'rh':
                        perfil = Perfil.query.filter_by(nome="RH").first()
                        if not perfil:
                            perfil = Perfil(nome="RH", descricao="Recursos Humanos")
                            db.session.add(perfil)
                            db.session.commit()
                    elif novo_tipo == 'temporario':
                        perfil = Perfil.query.filter_by(nome="Temporário").first()
                        if not perfil:
                            perfil = Perfil(nome="Temporário", descricao="Acesso temporário por 24h", acesso_temporario=True)
                            db.session.add(perfil)
                            db.session.commit()
                    else:
                        # Para outros tipos (se, am, comercial*)
                        perfil = Perfil.query.filter_by(nome=novo_tipo.upper()).first()
                        if not perfil:
                            perfil = Perfil(nome=novo_tipo.upper())
                            db.session.add(perfil)
                            db.session.commit()
                    
                    # Atualizar tipo e superusuário
                    usuario.id_perfil = perfil.id
                    usuario.superusuario = is_superusuario
                    usuario.data_expiracao = data_expiracao
                    
                    # Atualizar permissões de blocos
                    usuario.blocos_permitidos = []
                    for nome_bloco in blocos_permitidos:
                        bloco = BlocoProposta.query.filter_by(nome=nome_bloco).first()
                        if bloco:
                            usuario.blocos_permitidos.append(bloco)
                    
                    db.session.commit()
                    
                    # Backup em JSON
                    usuarios = carregar_usuarios()
                    if usuario_alterar in usuarios:
                        usuarios[usuario_alterar]['tipo'] = novo_tipo
                        usuarios[usuario_alterar]['superusuario'] = is_superusuario
                        if data_expiracao:
                            usuarios[usuario_alterar]['data_expiracao'] = data_expiracao.strftime('%d/%m/%Y %H:%M:%S')
                        salvar_usuarios(usuarios)
                
                flash('Tipo de usuário alterado com sucesso!', 'success')
            except Exception as e:
                flash(f'Erro ao alterar tipo do usuário: {str(e)}', 'danger')
                logging.error(f"Erro ao alterar tipo do usuário: {e}")
    
    # Carregar ofertas para incluir seus blocos nos seletores
    ofertas = {}
    blocos_por_categoria = {"geral": {}}
    
    try:
        # Carregar ofertas
        ofertas = carregar_ofertas()
        logging.info(f"Carregadas {len(ofertas)} ofertas para permissões de acesso")
        
        # Organizar blocos por categoria para melhor visualização
        for nome_bloco, bloco in blocos_db.items():
            categoria = bloco.get("categoria", "geral")
            
            if categoria not in blocos_por_categoria:
                blocos_por_categoria[categoria] = {}
            
            blocos_por_categoria[categoria][nome_bloco] = bloco
            
            # Adicionar também à categoria "geral" para visualização completa
            if categoria != "geral":
                blocos_por_categoria["geral"][nome_bloco] = bloco
                
        # Processar blocos específicos de ofertas
        for oferta_nome, oferta_dados in ofertas.items():
            categoria_oferta = f"oferta_{oferta_nome}"
            
            # Garantir que a categoria existe
            if categoria_oferta not in blocos_por_categoria:
                blocos_por_categoria[categoria_oferta] = {}
            
            # Processar blocos da oferta se existirem
            for bloco_nome, bloco_info in oferta_dados.get('blocos', {}).items():
                # Verificar se o bloco já existe
                if bloco_nome not in blocos_db:
                    # Criar o bloco temporariamente para exibição
                    is_obrigatorio = bloco_nome in oferta_dados.get('obrigatorios', [])
                    novo_bloco = {
                        "titulo": bloco_info.get('titulo', bloco_nome.replace('_', ' ')),
                        "texto": bloco_info.get('texto', ''),
                        "obrigatorio": is_obrigatorio,
                        "categoria": categoria_oferta
                    }
                    blocos_db[bloco_nome] = novo_bloco
                    blocos_por_categoria[categoria_oferta][bloco_nome] = novo_bloco
    except Exception as e:
        logging.error(f"Erro ao processar blocos de ofertas: {e}")
    
    return render_template('gerenciar_usuarios.html', 
                          usuarios=usuarios_carregados, 
                          perfis=perfis_disponiveis, 
                          blocos=blocos_db,
                          blocos_por_categoria=blocos_por_categoria,
                          ofertas=ofertas)

# Rota para adicionar usuário (mantida para compatibilidade)
@app.route('/adicionar_usuario', methods=['POST'])
@admin_required
def adicionar_usuario():
    try:
        nome = request.form.get('nome', '')
        login = request.form.get('login', '')
        senha = request.form.get('senha', '')
        perfil = request.form.get('perfil', '')
        tipo_acesso = request.form.get('tipo_acesso', 'todos')
        blocos_permitidos = request.form.getlist('blocos_permitidos')
        
        # Debug logs
        logging.info(f"Dados recebidos para adicionar usuário: nome={nome}, login={login}, perfil={perfil}")
        logging.info(f"Tipo de acesso selecionado: {tipo_acesso}")
        logging.info(f"Blocos permitidos: {blocos_permitidos}")
        logging.info(f"Total de blocos selecionados: {len(blocos_permitidos)}")
        
        if not nome or not login or not senha or not perfil:
            flash('Todos os campos são obrigatórios!', 'danger')
            return redirect(url_for('gerenciar_usuarios'))
    
        # Verificar se o login já existe usando SQL nativo
        from sqlalchemy import text
        check_query = text("SELECT COUNT(*) FROM usuario WHERE login = :login")
        result = db.session.execute(check_query, {"login": login}).scalar()
    
        if result > 0:
            flash('Este login já está em uso!', 'danger')
            return redirect(url_for('gerenciar_usuarios'))
    
        # Obter ID do perfil selecionado
        perfil_query = text("SELECT id FROM perfil WHERE nome = :nome")
        perfil_id = db.session.execute(perfil_query, {"nome": perfil}).scalar()
        
        if not perfil_id:
            # Se o perfil não existir, vamos criar sem usar a coluna acesso_temporario
            perfil_insert = text("""
                INSERT INTO perfil (nome) 
                VALUES (:nome)
                RETURNING id
            """)
            perfil_id = db.session.execute(perfil_insert, {"nome": perfil}).scalar()
            
        # Calcular data de expiração para usuários temporários (24 horas)
        data_expiracao = None
        if perfil == "Temporário":
            data_expiracao = datetime.datetime.utcnow() + datetime.timedelta(hours=24)
            
        # Determinar se é superusuário baseado no tipo de acesso
        is_superuser = tipo_acesso == 'todos'
        
        logging.info(f"Criando usuário com perfil_id={perfil_id}, superusuario={is_superuser}")
        
        # Inserir o usuário usando SQL Nativo para evitar problemas com colunas faltantes
        try:
            # Construir consulta básica que funciona para qualquer banco
            insert_query = text("""
                INSERT INTO usuario (nome, login, senha, status, id_perfil, superusuario)
                VALUES (:nome, :login, :senha, :status, :id_perfil, :superusuario)
                RETURNING id
            """)
            
            params = {
                "nome": nome,
                "login": login,
                "senha": generate_password_hash(senha),
                "status": 1,
                "id_perfil": perfil_id,
                "superusuario": is_superuser
            }
            
            usuario_id = db.session.execute(insert_query, params).scalar()
            logging.info(f"Usuário criado com ID: {usuario_id}")
                
        except Exception as e:
            logging.error(f"Erro ao inserir usuário: {e}")
            # Tentar versão mais simples caso a versão com superusuario falhe
            try:
                basic_query = text("""
                    INSERT INTO usuario (nome, login, senha, status, id_perfil)
                    VALUES (:nome, :login, :senha, :status, :id_perfil)
                    RETURNING id
                """)
                
                params = {
                    "nome": nome,
                    "login": login,
                    "senha": generate_password_hash(senha),
                    "status": 1,
                    "id_perfil": perfil_id
                }
                
                usuario_id = db.session.execute(basic_query, params).scalar()
                logging.info(f"Usuário criado com versão simplificada, ID: {usuario_id}")
            except Exception as e2:
                logging.error(f"Erro na tentativa simplificada: {e2}")
                raise e2
        
        # Se o tipo de acesso for específico e tiver blocos selecionados
        if tipo_acesso == 'especificos' and blocos_permitidos and usuario_id:
            logging.info(f"Processando {len(blocos_permitidos)} blocos permitidos para o usuário {login}")
            
            # Para cada bloco na lista de blocos permitidos
            for bloco_nome in blocos_permitidos:
                try:
                    # Primeiro verificar se o bloco existe (usando SQL nativo)
                    check_bloco = text("SELECT COUNT(*) FROM bloco_proposta WHERE nome = :nome")
                    bloco_existe = db.session.execute(check_bloco, {"nome": bloco_nome}).scalar() > 0
                    
                    if bloco_existe:
                        # Verificar se a associação já existe
                        check_perm = text("""
                            SELECT COUNT(*) FROM usuario_permissoes_blocos 
                            WHERE usuario_id = :usuario_id AND bloco_nome = :bloco_nome
                        """)
                        perm_existe = db.session.execute(check_perm, {
                            "usuario_id": usuario_id, 
                            "bloco_nome": bloco_nome
                        }).scalar() > 0
                        
                        if not perm_existe:
                            # Adicionar permissão
                            add_perm = text("""
                                INSERT INTO usuario_permissoes_blocos (usuario_id, bloco_nome)
                                VALUES (:usuario_id, :bloco_nome)
                            """)
                            db.session.execute(add_perm, {
                                "usuario_id": usuario_id,
                                "bloco_nome": bloco_nome
                            })
                            logging.info(f"Adicionada permissão para o bloco: {bloco_nome}")
                    else:
                        # Se o bloco não existir, tente inseri-lo
                        logging.warning(f"Bloco não encontrado no banco de dados: {bloco_nome}, tentando criar...")
                        novo_bloco = BlocoProposta(
                            nome=bloco_nome,
                            titulo=bloco_nome.replace('_', ' '),
                            texto="",
                            obrigatorio=False,
                            criado_por="sistema"
                        )
                        db.session.add(novo_bloco)
                        db.session.flush()
                        
                        # Adicionar permissão para o novo bloco
                        add_perm = text("""
                            INSERT INTO usuario_permissoes_blocos (usuario_id, bloco_nome)
                            VALUES (:usuario_id, :bloco_nome)
                        """)
                        db.session.execute(add_perm, {
                            "usuario_id": usuario_id,
                            "bloco_nome": bloco_nome
                        })
                        logging.info(f"Criado novo bloco e adicionada permissão: {bloco_nome}")
                except Exception as e:
                    logging.error(f"Erro ao processar bloco {bloco_nome}: {e}")
        
        # Commit das alterações
        db.session.commit()
        
        # Atualizar o arquivo JSON como backup
        usuarios = carregar_usuarios()
        usuarios[login] = {
            "nome": nome,
            "senha": generate_password_hash(senha),
            "status": 1,
            "tipo": "admin" if perfil == "Governança" else perfil.lower(),
            "perfil": perfil,
            "superusuario": is_superuser
        }
        salvar_usuarios(usuarios)
        
        flash(f'Usuário {nome} adicionado com sucesso!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao adicionar usuário: {str(e)}', 'danger')
        logging.error(f"Erro ao adicionar usuário: {e}")
    
    return redirect(url_for('gerenciar_usuarios'))

# Função para salvar proposta como rascunho
def salvar_como_rascunho(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados, rascunho_id=None, logo_atual=None, oferta_selecionada=None, blocos_temporarios=None):
    # Verificar se há dados suficientes para salvar
    if not nome_cliente:
        flash('Por favor, informe pelo menos o nome do cliente para salvar como rascunho.')
        return redirect(url_for('exibir_criar_proposta'))
    
    # Carregar rascunhos existentes
    rascunhos = carregar_rascunhos()
    
    # Gerar ID para o rascunho se não existir
    if not rascunho_id:
        rascunho_id = str(uuid.uuid4())
    
    # Processar upload de arquivo de logo
    if 'logo_file' in request.files and request.files['logo_file'].filename:
        logo_file = request.files['logo_file']
        # Verificar extensão
        if logo_file and allowed_file(logo_file.filename):
            # Gerar nome único para o arquivo
            filename = secure_filename(f"{nome_cliente.replace(' ', '_')}_{uuid.uuid4()}.{logo_file.filename.rsplit('.', 1)[1].lower()}")
            logo_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            logo_file.save(logo_path)
            # Salvar o caminho relativo para facilitar a portabilidade
            logo_cliente = os.path.join('static', 'uploads', filename)
            logging.info(f"Logo salvo em: {logo_path}, caminho relativo: {logo_cliente}")
        else:
            flash('Formato de arquivo não suportado. Use PNG, JPG ou GIF.')
            logger.warning(f"Formato de arquivo não suportado: {logo_file.filename}")
            return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
    elif logo_atual:
        # Se não foi enviado novo arquivo, mas existe um logo atual, usar o logo atual
        logo_cliente = logo_atual
        logging.info(f"Usando logo atual: {logo_cliente}")
    
    # Verificar se o logo existe
    if logo_cliente:
        # Verificar se o caminho é absoluto ou relativo
        logo_path = logo_cliente
        if not os.path.isabs(logo_cliente):
            logo_path = os.path.join(app.root_path, logo_cliente)
        
        if not os.path.exists(logo_path):
            logging.warning(f"Logo não encontrado no caminho: {logo_path}")
            logo_cliente = None
    
    # Garantir que blocos_selecionados seja uma lista
    if not isinstance(blocos_selecionados, list):
        blocos_selecionados = [blocos_selecionados] if blocos_selecionados else []
    
    # Garantir que blocos_temporarios seja um dicionário
    if not blocos_temporarios:
        blocos_temporarios = {}
    elif not isinstance(blocos_temporarios, dict):
        blocos_temporarios = {}
    
    # Criar ou atualizar o rascunho
    rascunhos[rascunho_id] = {
        'nome_cliente': nome_cliente,
        'logo_cliente': logo_cliente,
        'blocos_selecionados': blocos_selecionados,
        'usuario': session.get('usuario_logado'),
        'data_atualizacao': datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
        'oferta_selecionada': oferta_selecionada,
        'blocos_temporarios': blocos_temporarios  # Adicionar blocos temporários ao rascunho
    }
    
    # Salvar os rascunhos
    salvar_rascunhos(rascunhos)
    
    # Verificar se é um salvamento automático
    auto_save = request.form.get('auto_save', '')
    if auto_save == '1':
        # Para salvamento automático, retornar uma resposta vazia com status 204 (No Content)
        return ('', 204)
    
    flash('Rascunho salvo com sucesso!')
    return redirect(url_for('dashboard'))

@app.route('/excluir_rascunho/<rascunho_id>')
@login_required
def excluir_rascunho(rascunho_id):
    # Carregar rascunhos existentes
    rascunhos = carregar_rascunhos()
    if rascunho_id in rascunhos:
        # Verificar se o usuário é admin ou o criador do rascunho
        if session.get('tipo_usuario') == 'admin' or rascunhos[rascunho_id].get('usuario') == session.get('usuario_logado'):
            # Remover o rascunho do banco de dados
            try:
                rascunho_db = Rascunho.query.get(rascunho_id)
                if rascunho_db:
                    db.session.delete(rascunho_db)
                    db.session.commit()
                    app.logger.info(f"Rascunho {rascunho_id} removido do banco de dados")
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"Erro ao remover rascunho do banco de dados: {str(e)}")
            
            # Remover o logo do rascunho se existir
            logo_path = rascunhos[rascunho_id].get('logo_cliente', '')
            if logo_path and os.path.exists(logo_path) and 'placeholder-logo' not in logo_path:
                try:
                    os.remove(logo_path)
                except Exception as e:
                    app.logger.error(f"Erro ao remover logo do rascunho: {str(e)}")
            
            # Remover o rascunho do dicionário
            del rascunhos[rascunho_id]
            # Salvar as alterações
            salvar_rascunhos(rascunhos)
            
            flash('Rascunho excluído com sucesso.')
        else:
            flash('Você não tem permissão para excluir este rascunho.')
    else:
        flash('Rascunho não encontrado.')
    
    # Redirecionar para a página anterior
    previous_page = request.referrer
    if previous_page and url_for('exibir_criar_proposta') in previous_page:
        return redirect(previous_page)
    return redirect(url_for('dashboard'))

@app.route('/static/img/placeholder-logo.png')
def placeholder_logo():
    # Criar uma imagem de placeholder
    width, height = 200, 200
    image = Image.new('RGB', (width, height), color=(240, 240, 240))
    
    # Adicionar texto
    draw = ImageDraw.Draw(image)
    
    # Tentar carregar uma fonte, ou usar a fonte padrão
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except IOError:
        font = ImageFont.load_default()
    
    text = "Logo do Cliente"
    
    # Diferentes versões do PIL têm diferentes métodos para obter o tamanho do texto
    try:
        # Para versões mais recentes do PIL
        left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
        text_width = right - left
        text_height = bottom - top
    except AttributeError:
        try:
            # Para versões intermediárias do PIL
            text_width, text_height = draw.textsize(text, font=font)
        except AttributeError:
            # Fallback para valores aproximados
            text_width, text_height = 100, 20
    
    position = ((width - text_width) // 2, (height - text_height) // 2)
    
    # Desenhar o texto
    draw.text(position, text, fill=(150, 150, 150), font=font)
    
    # Salvar a imagem em um buffer
    img_io = BytesIO()
    image.save(img_io, 'PNG')
    img_io.seek(0)
    
    return send_file(img_io, mimetype='image/png')

# Função para carregar ofertas do arquivo JSON
def carregar_ofertas():
    try:
        if os.path.exists(OFERTAS_FILE):
            with open(OFERTAS_FILE, "r", encoding="utf-8") as f:
                ofertas = json.load(f)
                return ofertas
        else:
            return {}
    except Exception as e:
        app.logger.error(f"Erro ao carregar ofertas: {e}")
        return {}

# Função para salvar ofertas no arquivo JSON
def salvar_ofertas(ofertas):
    try:
        # Salvar no banco de dados
        for tipo_oferta in ofertas.keys():
            # Verificar se oferta já existe
            oferta = Oferta.query.filter_by(tipo_oferta=tipo_oferta).first()
            
            if not oferta:
                # Criar nova oferta
                nova_oferta = Oferta(tipo_oferta=tipo_oferta)
                db.session.add(nova_oferta)
        
        db.session.commit()
    
        # Backup em JSON
        with open(OFERTAS_FILE, "w", encoding="utf-8") as f:
            json.dump(ofertas, f, ensure_ascii=False, indent=4)
            
    except Exception as e:
        app.logger.error(f"Erro ao salvar ofertas no banco: {e}")
        # Fallback para JSON
        with open(OFERTAS_FILE, "w", encoding="utf-8") as f:
            json.dump(ofertas, f, ensure_ascii=False, indent=4)

# Rota para obter todos os blocos (API)
@app.route('/api/blocos', methods=['GET'])
@login_required
def api_blocos():
    try:
        blocos = carregar_blocos()
        return jsonify(blocos)
    except Exception as e:
        app.logger.error(f"Erro ao carregar blocos: {str(e)}")
        return jsonify({"error": "Erro ao carregar blocos"}), 500

# Rota para obter um bloco específico (API)
@app.route('/api/bloco/<bloco_nome>', methods=['GET'])
@login_required
def api_bloco(bloco_nome):
    try:
        app.logger.info(f"Solicitação para obter bloco: {bloco_nome}")
        blocos = carregar_blocos()
        app.logger.info(f"Blocos carregados: {list(blocos.keys())}")
        
        if bloco_nome in blocos:
            app.logger.info(f"Bloco encontrado: {bloco_nome}")
            return jsonify(blocos[bloco_nome])
        else:
            app.logger.warning(f"Bloco não encontrado: {bloco_nome}")
            # Retornar um template vazio para permitir a criação do bloco
            return jsonify({
                "error": "Bloco não encontrado", 
                "texto": "<p>Este é um novo bloco. Edite o conteúdo e salve para criá-lo.</p>",
                "novo_bloco": True
            }), 200  # Retornar 200 em vez de 404 para permitir a criação
    except Exception as e:
        app.logger.error(f"Erro ao carregar bloco {bloco_nome}: {str(e)}")
        return jsonify({"error": "Erro ao carregar bloco", "texto": "<p>Erro ao carregar bloco</p>"}), 500

# Rota para salvar um bloco (API)
@app.route('/api/salvar_bloco', methods=['POST'])
@login_required
def salvar_bloco_api():
    try:
        data = request.json
        bloco_nome = data.get('nome')
        texto = data.get('texto')
        cliente = data.get('cliente')
        rascunho_id = data.get('rascunho_id')
        temporario = data.get('temporario', False)
        
        if not bloco_nome or not texto:
            return jsonify({"success": False, "error": "Dados incompletos"}), 400
        
        # Se for temporário e tiver um rascunho_id, salvar no rascunho
        if temporario and rascunho_id:
            rascunhos = carregar_rascunhos()
            
            if rascunho_id in rascunhos:
                rascunho = rascunhos[rascunho_id]
                
                # Adicionar como bloco temporário no rascunho
                if 'blocos_temporarios' not in rascunho:
                    rascunho['blocos_temporarios'] = {}
                
                # Adicionar o bloco temporário
                bloco_dados = {
                    'texto': texto,
                    'criado_por': session.get('usuario_logado'),
                    'data_criacao': datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')
                }
                
                rascunho['blocos_temporarios'][bloco_nome] = bloco_dados
                
                # Adicionar o bloco à lista de blocos selecionados do rascunho
                if 'blocos_selecionados' not in rascunho:
                    rascunho['blocos_selecionados'] = []
                
                if bloco_nome not in rascunho['blocos_selecionados']:
                    rascunho['blocos_selecionados'].append(bloco_nome)
                
                # Salvar o rascunho atualizado
                rascunhos[rascunho_id] = rascunho
                salvar_rascunhos(rascunhos)
                logging.info(f"Bloco temporário adicionado ao rascunho: {bloco_nome}")
                return jsonify({
                    "success": True,
                    "bloco": {
                        "nome": bloco_nome,
                        "texto": texto,
                        "temporario": True,
                        "cliente": cliente
                    }
                })
            else:
                logging.warning(f"Rascunho não encontrado: {rascunho_id}")
                # Se o rascunho não existir, criar um bloco permanente mesmo assim
                logging.info(f"Criando bloco permanente em vez de temporário: {bloco_nome}")
        
        # Caso contrário, salvar como bloco permanente
        blocos = carregar_blocos()
        
        if bloco_nome in blocos:
            # Atualizar bloco existente
            blocos[bloco_nome]['texto'] = texto
            logging.info(f"Bloco atualizado: {bloco_nome}")
        else:
            # Criar novo bloco
            blocos[bloco_nome] = {
                'texto': texto,
                'imagem': None,
                'obrigatorio': False,
                'criado_por': session.get('usuario_logado'),
                'data_criacao': datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                'cliente_associado': cliente,
                'reutilizavel': True
            }
            logging.info(f"Novo bloco criado: {bloco_nome}")
        
        salvar_blocos(blocos)
        return jsonify({
            "success": True,
            "bloco": {
                "nome": bloco_nome,
                "texto": texto,
                "temporario": False,
                "cliente": cliente
            }
        })
    except Exception as e:
        logging.error(f"Erro ao salvar bloco: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

# Adicionar nova rota para excluir bloco
@app.route('/api/excluir_bloco/<bloco_id>', methods=['DELETE'])
@login_required
def excluir_bloco_api(bloco_id):
    try:
        blocos = carregar_blocos()
        if bloco_id in blocos:
            # Verificar se o usuário é admin
            if session.get('tipo_usuario') != 'admin':
                return jsonify({"success": False, "error": "Apenas administradores podem excluir blocos"}), 403
                
            del blocos[bloco_id]
            salvar_blocos(blocos)
            return jsonify({"success": True})
        return jsonify({"success": False, "error": "Bloco não encontrado"}), 404
    except Exception as e:
        app.logger.error(f"Erro ao excluir bloco: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/adicionar_bloco', methods=['GET', 'POST'])
@login_required
def adicionar_bloco():
    cliente = request.args.get('cliente', '')
    
    if request.method == 'POST':
        nome_bloco = request.form['nome_bloco'].replace(' ', '_')
        texto_bloco = request.form['texto_bloco']
        cliente_associado = request.form.get('cliente_associado', '')
        reutilizavel = request.form.get('reutilizavel', 'off') == 'on'
        
        blocos = carregar_blocos()
        
        # Verificar se o bloco já existe
        if nome_bloco in blocos and reutilizavel:
            flash('Já existe um bloco com este nome. Escolha outro nome para o bloco reutilizável.')
            return render_template('adicionar_bloco.html', cliente=cliente)
        
        # Se não for reutilizável, tornar o nome do bloco único para este cliente
        if not reutilizavel and cliente_associado:
            nome_bloco = f"{cliente_associado}_{nome_bloco}"
        
        # Adicionar o novo bloco
        blocos[nome_bloco] = {
            'texto': texto_bloco,
            'imagem': None,
            'obrigatorio': False,
            'criado_por': session.get('usuario_logado'),
            'data_criacao': datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            'cliente_associado': cliente_associado,
            'reutilizavel': reutilizavel
        }
        
        salvar_blocos(blocos)
        flash('Bloco adicionado com sucesso!')
        
        # Redirecionar para a página de criação de proposta, passando o bloco como selecionado
        if cliente_associado:
            # Se estamos criando um bloco para uma proposta específica, pré-selecionar o bloco
            return redirect(url_for('exibir_criar_proposta', cliente=cliente_associado, bloco_novo=nome_bloco))
        else:
            return redirect(url_for('exibir_criar_proposta'))
    
    return render_template('adicionar_bloco.html', cliente=cliente)

# Rota para servir o favicon.ico
@app.route('/favicon.ico')
def favicon():
    return app.send_static_file('favicon_round.png')

# Rota para servir o webmanifest
@app.route('/site.webmanifest')
def webmanifest():
    manifest = {
        "name": "Gerador de Propostas Service IT",
        "short_name": "Propostas SVIT",
        "icons": [
            {
                "src": "/static/favicon_round.png",
                "sizes": "192x192",
                "type": "image/png"
            }
        ],
        "theme_color": "#e60000",
        "background_color": "#ffffff",
        "display": "standalone"
    }
    return jsonify(manifest)


# Aplicar decorator às funções que usam o banco de dados
carregar_usuarios = with_app_context(carregar_usuarios)
salvar_usuarios = with_app_context(salvar_usuarios)
carregar_propostas = with_app_context(carregar_propostas)
salvar_propostas = with_app_context(salvar_propostas)
carregar_blocos = with_app_context(carregar_blocos)
salvar_blocos = with_app_context(salvar_blocos)
carregar_rascunhos = with_app_context(carregar_rascunhos)
salvar_rascunhos = with_app_context(salvar_rascunhos)
carregar_ofertas = with_app_context(carregar_ofertas)
salvar_ofertas = with_app_context(salvar_ofertas)

@app.route('/alterar_senha', methods=['GET', 'POST'])
@login_required
def alterar_senha():
    if request.method == 'POST':
        senha_atual = request.form.get('senha_atual')
        nova_senha = request.form.get('nova_senha')
        confirmar_senha = request.form.get('confirmar_senha')
        
        if not senha_atual or not nova_senha or not confirmar_senha:
            flash('Todos os campos são obrigatórios.', 'danger')
            return redirect(url_for('alterar_senha'))
        
        if nova_senha != confirmar_senha:
            flash('A nova senha e a confirmação não correspondem.', 'danger')
            return redirect(url_for('alterar_senha'))
        
        try:
            # Verificar credenciais atuais
            with app.app_context():
                usuario = Usuario.query.filter_by(login=session['usuario_logado']).first()
                
                if not usuario:
                    flash('Usuário não encontrado no sistema.', 'danger')
                    return redirect(url_for('dashboard'))
                
                # Verificar senha atual
                senha_valida = False
                if usuario.senha.startswith('pbkdf2:sha256:'):
                    senha_valida = check_password_hash(usuario.senha, senha_atual)
                else:
                    # Compatibilidade com senhas em texto puro
                    senha_valida = usuario.senha == senha_atual
                
                if not senha_valida:
                    flash('Senha atual incorreta.', 'danger')
                    return redirect(url_for('alterar_senha'))
                
                # Atualizar senha no banco de dados (sempre com hash)
                usuario.senha = generate_password_hash(nova_senha)
                db.session.commit()
                
                # Atualizar também no JSON para compatibilidade
                usuarios = carregar_usuarios()
                if session['usuario_logado'] in usuarios:
                    usuarios[session['usuario_logado']]['senha'] = nova_senha
                    salvar_usuarios(usuarios)
                
                flash('Senha alterada com sucesso!', 'success')
                return redirect(url_for('dashboard'))
        except Exception as e:
            app.logger.error(f"Erro ao alterar senha: {e}")
            flash('Ocorreu um erro ao alterar a senha. Tente novamente.', 'danger')
            return redirect(url_for('alterar_senha'))
    
    return render_template('alterar_senha.html')

@app.route('/obter_conteudo_bloco', methods=['GET'])
@login_required
def obter_conteudo_bloco():
    nome_bloco = request.args.get('nome')
    if not nome_bloco:
        return jsonify({'error': 'Nome do bloco não especificado'}), 400
        
    blocos = carregar_blocos()
    if nome_bloco in blocos:
        return jsonify({
            'conteudo': blocos[nome_bloco].get('texto', '')
        })
    return jsonify({'error': 'Bloco não encontrado'}), 404

def substituir_capa(document, nome_cliente, logo_cliente_path):
    """
    Função para substituir o marcador {{nome_cliente}} e {{logo_cliente}} na capa do documento
    com tratamento avançado para melhor apresentação visual
    """
    logging.info("Substituindo variáveis da capa com tratamento visual aprimorado")
    
    # Substituir NOME_CLIENTE e logo_cliente
    for paragraph in document.paragraphs:
        substituir_variaveis(paragraph, "NOME_CLIENTE", nome_cliente)
        
        # Inserir logotipo se o marcador for encontrado
        if ( ("{{logo_cliente}}" in paragraph.text or "logo_cliente" in paragraph.text) and logo_cliente_path):
            # Limpar completamente o parágrafo para evitar sobreposição
            paragraph.clear()
            # Criar um parágrafo específico para a logo e centralizá-lo
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            
            # Determinar o tamanho apropriado para a logo
            try:
                from PIL import Image
                img = Image.open(logo_cliente_path)
                width, height = img.size
                
                # Usar tamanhos mais conservadores para a logo
                max_width = Inches(2.0)  # Tamanho base menor que antes
                
                # Ajustar com base na proporção de forma mais conservadora
                if width > height * 2.5:  # Logo extremamente largo
                    max_width = Inches(2.2)
                elif width > height * 1.8:  # Logo muito largo
                    max_width = Inches(2.0)
                elif width > height * 1.2:  # Logo moderadamente largo
                    max_width = Inches(1.8)
                elif height > width * 1.8:  # Logo muito alto
                    max_width = Inches(1.2)
                elif height > width * 1.2:  # Logo moderadamente alto
                    max_width = Inches(1.5)
                    
                # Garantir que a altura não fique excessiva - limite mais restritivo
                height_calculated = max_width * height / width if width > 0 else Inches(1.2)
                max_height = Inches(1.2)  # Limite máximo de altura reduzido
                
                if height_calculated > max_height:
                    max_width = max_width * max_height / height_calculated
                
                # Aprimorar o espaçamento para melhor apresentação visual
                paragraph.space_before = Pt(24)  # Mais espaço acima
                paragraph.space_after = Pt(36)   # Mais espaço abaixo
                
                # Configurar margens para garantir que não haja problemas de posicionamento
                for section in document.sections:
                    if section.start_type == WD_SECTION_START.NEW_PAGE:
                        section.top_margin = Inches(1.2)
                        section.left_margin = Inches(1.0)
                        section.right_margin = Inches(1.0)
                
                # Adicionar a logo com o tamanho calculado e tratamento de qualidade
                run.add_picture(logo_cliente_path, width=max_width)
                logging.info(f"Logo adicionado à capa com dimensionamento otimizado: largura={max_width}")
            except Exception as e:
                # Fallback para tamanho padrão em caso de erro, mas ainda com bom dimensionamento
                logging.error(f"Erro ao calcular dimensões do logo: {e}, usando tamanho padrão aprimorado")
                run.add_picture(logo_cliente_path, width=Inches(1.8))
                paragraph.space_before = Pt(24)
                paragraph.space_after = Pt(36)
    
    # Substituir também no cabeçalho e rodapé com tratamento visual
    for section in document.sections:
        header = section.header
        for paragraph in header.paragraphs:
            substituir_variaveis(paragraph, "NOME_CLIENTE", nome_cliente)
            
            # Inserir logotipo se o marcador for encontrado
            if ( ("{{logo_cliente}}" in paragraph.text or "logo_cliente" in paragraph.text) and logo_cliente_path):
                # Limpar completamente o parágrafo
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                
                # Usar um tamanho menor para cabeçalhos
                try:
                    from PIL import Image
                    img = Image.open(logo_cliente_path)
                    width, height = img.size
                    
                    # Tamanho menor para cabeçalhos
                    max_width = Inches(0.8) 
                    
                    # Ajuste com base na proporção, mas mantendo dimensões reduzidas
                    ratio = width / height
                    if ratio > 2:  # Logo muito largo
                        max_width = Inches(1.0)
                    elif ratio < 0.5:  # Logo muito alto
                        max_width = Inches(0.6)
                    
                    # Adicionar espaço ao redor da logo no cabeçalho
                    paragraph.space_before = Pt(6)
                    paragraph.space_after = Pt(6)
                    
                    run.add_picture(logo_cliente_path, width=max_width)
                    logging.info(f"Logo adicionado ao cabeçalho com tamanho adaptativo: largura={max_width}")
                except Exception as e:
                    logging.error(f"Erro ao calcular dimensões do logo para cabeçalho: {e}, usando tamanho padrão")
                    run.add_picture(logo_cliente_path, width=Inches(0.8))
    
    # Substituir também no rodapé com tratamento visual
    for section in document.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            substituir_variaveis(paragraph, "NOME_CLIENTE", nome_cliente)

# Rota para gerenciar blocos e suas permissões
@app.route('/gerenciar_blocos', methods=['GET', 'POST'])
@login_required
def gerenciar_blocos():
    # Verificar se o método é POST e o usuário não é admin para bloquear ações não permitidas
    if request.method == 'POST' and not session.get('tipo_usuario') == 'admin':
        flash('Você não tem permissão para realizar esta ação.', 'danger')
        return redirect(url_for('gerenciar_blocos'))
    
    if request.method == 'POST':
        acao = request.form.get('acao')
        
        if acao == 'adicionar':
            # Adicionar novo bloco (apenas admin)
            nome_bloco = request.form.get('nome_bloco')
            titulo_bloco = request.form.get('titulo_bloco')
            conteudo_bloco = request.form.get('conteudo_bloco', '')
            obrigatorio = request.form.get('obrigatorio') == 'on'
            usuarios_permitidos = request.form.getlist('usuarios_permitidos')
            
            if not nome_bloco or not titulo_bloco:
                flash('Nome e título do bloco são obrigatórios.', 'danger')
                return redirect(url_for('gerenciar_blocos'))
            
            try:
                # Verificar se o bloco já existe
                bloco_existente = BlocoProposta.query.filter_by(nome=nome_bloco).first()
                if bloco_existente:
                    flash('Já existe um bloco com este nome.', 'danger')
                    return redirect(url_for('gerenciar_blocos'))
                
                # Criar o novo bloco
                novo_bloco = BlocoProposta(
                    nome=nome_bloco,
                    titulo=titulo_bloco,
                    texto=conteudo_bloco,
                    obrigatorio=obrigatorio,
                    criado_por=session.get('usuario_logado'),
                    data_criacao=datetime.datetime.utcnow()
                )
                db.session.add(novo_bloco)
                db.session.flush()  # Para obter o ID do bloco antes de commit
                
                # Adicionar permissões de usuários
                for login_usuario in usuarios_permitidos:
                    usuario = Usuario.query.filter_by(login=login_usuario).first()
                    if usuario:
                        usuario.blocos_permitidos.append(novo_bloco)
                
                db.session.commit()
                
                # Backup em JSON
                blocos = carregar_blocos()
                blocos[nome_bloco] = {
                    "texto": conteudo_bloco,
                    "titulo": titulo_bloco,
                    "obrigatorio": obrigatorio,
                    "criado_por": session.get('usuario_logado'),
                    "data_criacao": datetime.datetime.utcnow().strftime('%d/%m/%Y %H:%M:%S')
                }
                salvar_blocos(blocos)
                
                flash('Bloco adicionado com sucesso!', 'success')
            except Exception as e:
                flash(f'Erro ao adicionar bloco: {str(e)}', 'danger')
                logging.error(f"Erro ao adicionar bloco: {e}")
        
        elif acao == 'editar':
            # Editar bloco existente
            nome_bloco = request.form.get('nome_bloco')
            titulo_bloco = request.form.get('titulo_bloco')
            conteudo_bloco = request.form.get('conteudo_bloco', '')
            obrigatorio = request.form.get('obrigatorio') == 'on'
            usuarios_permitidos = request.form.getlist('usuarios_permitidos')
            
            if not nome_bloco:
                flash('Nome do bloco é obrigatório.', 'danger')
                return redirect(url_for('gerenciar_blocos'))
            
            try:
                # Buscar o bloco
                bloco = BlocoProposta.query.filter_by(nome=nome_bloco).first()
                if bloco:
                    # Validar permissão (apenas admin ou o criador do bloco podem editar)
                    usuario_atual = session.get('usuario_logado')
                    if session.get('tipo_usuario') != 'admin' and bloco.criado_por != usuario_atual:
                        flash('Você não tem permissão para editar este bloco.', 'danger')
                        return redirect(url_for('gerenciar_blocos'))
                    
                    # Atualizar dados do bloco
                    bloco.titulo = titulo_bloco
                    bloco.texto = conteudo_bloco
                    
                    # Apenas admins podem alterar se o bloco é obrigatório
                    if session.get('tipo_usuario') == 'admin':
                        bloco.obrigatorio = obrigatorio
                        
                        # Limpar as permissões atuais
                        for usuario in Usuario.query.all():
                            if bloco in usuario.blocos_permitidos:
                                usuario.blocos_permitidos.remove(bloco)
                        
                        # Adicionar novas permissões
                        for login_usuario in usuarios_permitidos:
                            usuario = Usuario.query.filter_by(login=login_usuario).first()
                            if usuario:
                                usuario.blocos_permitidos.append(bloco)
                    
                    db.session.commit()
                    
                    # Backup em JSON
                    blocos = carregar_blocos()
                    if nome_bloco in blocos:
                        blocos[nome_bloco]["titulo"] = titulo_bloco
                        blocos[nome_bloco]["texto"] = conteudo_bloco
                        if session.get('tipo_usuario') == 'admin':
                            blocos[nome_bloco]["obrigatorio"] = obrigatorio
                    salvar_blocos(blocos)
                    
                    flash('Bloco atualizado com sucesso!', 'success')
                else:
                    flash('Bloco não encontrado.', 'danger')
            except Exception as e:
                flash(f'Erro ao editar bloco: {str(e)}', 'danger')
                logging.error(f"Erro ao editar bloco: {e}")
        
        elif acao == 'remover':
            # Remover bloco (apenas admin)
            if session.get('tipo_usuario') != 'admin':
                flash('Apenas administradores podem remover blocos.', 'danger')
                return redirect(url_for('gerenciar_blocos'))
                
            nome_bloco = request.form.get('nome_bloco')
            
            if not nome_bloco:
                flash('Nome do bloco é obrigatório.', 'danger')
                return redirect(url_for('gerenciar_blocos'))
            
            try:
                # Verificar se é um bloco obrigatório padrão
                blocos_reservados = ["Termo_de_Confidencialidade", "Folha_de_Rosto", "Indice", "Introducao", "Sobre_a_ServiceIT"]
                if nome_bloco in blocos_reservados:
                    flash('Não é possível remover blocos reservados do sistema.', 'danger')
                    return redirect(url_for('gerenciar_blocos'))
                
                # Buscar o bloco
                bloco = BlocoProposta.query.filter_by(nome=nome_bloco).first()
                if bloco:
                    # Remover permissões relacionadas
                    for usuario in Usuario.query.all():
                        if bloco in usuario.blocos_permitidos:
                            usuario.blocos_permitidos.remove(bloco)
                    
                    # Remover o bloco
                    db.session.delete(bloco)
                    db.session.commit()
                    
                    # Backup em JSON
                    blocos = carregar_blocos()
                    if nome_bloco in blocos:
                        del blocos[nome_bloco]
                    salvar_blocos(blocos)
                    
                    flash('Bloco removido com sucesso!', 'success')
                else:
                    flash('Bloco não encontrado.', 'danger')
            except Exception as e:
                flash(f'Erro ao remover bloco: {str(e)}', 'danger')
                logging.error(f"Erro ao remover bloco: {e}")
    
    # Carregar blocos com base nas permissões do usuário
    blocos_db = {}
    try:
        usuario_atual = session.get('usuario_logado')
        tipo_usuario = session.get('tipo_usuario')
        is_admin = tipo_usuario == 'admin'
        
        # Obter o usuário do banco de dados
        usuario_obj = Usuario.query.filter_by(login=usuario_atual).first()
        is_superusuario = usuario_obj and usuario_obj.superusuario
        
        # Carregar do banco de dados
        query = BlocoProposta.query
        
        for bloco in query.all():
            # Verificar se o usuário tem permissão para este bloco
            tem_permissao = False
            
            # Admins e superusuários veem todos os blocos
            if is_admin or is_superusuario:
                tem_permissao = True
            else:
                # Verificar se o usuário tem permissão específica para este bloco
                if usuario_obj and bloco in usuario_obj.blocos_permitidos:
                    tem_permissao = True
                # Verificar se o usuário é o criador do bloco
                elif bloco.criado_por == usuario_atual:
                    tem_permissao = True
            
            # Só adicionar o bloco se o usuário tiver permissão
            if tem_permissao:
                # Obter usuários com permissão
                usuarios_permitidos = []
                for usuario in Usuario.query.all():
                    if bloco in usuario.blocos_permitidos or usuario.superusuario:
                        usuarios_permitidos.append(usuario.login)
                
                blocos_db[bloco.nome] = {
                    "titulo": bloco.titulo or bloco.nome.replace('_', ' '),
                    "texto": bloco.texto or "",
                    "obrigatorio": bloco.obrigatorio,
                    "criado_por": bloco.criado_por,
                    "data_criacao": bloco.data_criacao.strftime('%d/%m/%Y %H:%M:%S') if bloco.data_criacao else "",
                    "usuarios_permitidos": usuarios_permitidos
                }
    except Exception as e:
        logging.error(f"Erro ao carregar blocos do banco: {e}")
        # Fallback para JSON em caso de erro
        blocos_db = carregar_blocos()
    
    # Carregar usuários
    usuarios = {}
    try:
        for usuario in Usuario.query.all():
            perfil = Perfil.query.get(usuario.id_perfil)
            tipo_usuario = "usuario"
            if perfil:
                if perfil.nome == "Governança":
                    tipo_usuario = "admin"
                else:
                    tipo_usuario = perfil.nome.lower()
            
            usuarios[usuario.login] = {
                "nome": usuario.nome,
                "tipo": tipo_usuario,
                "status": usuario.status,
                "superusuario": usuario.superusuario
            }
    except Exception as e:
        logging.error(f"Erro ao carregar usuários do banco: {e}")
        # Fallback para JSON em caso de erro
        usuarios = carregar_usuarios()
    
    # Passar informação sobre o tipo de usuário para o template
    is_admin = session.get('tipo_usuario') == 'admin'
    
    return render_template('gerenciar_blocos.html', 
                           blocos=blocos_db, 
                           usuarios=usuarios, 
                           is_admin=is_admin)

# Função auxiliar para carregar blocos diretamente do arquivo JSON
def carregar_blocos_direto_json():
    try:
        blocos_json_path = os.path.join('data', 'blocos.json')
        if os.path.exists(blocos_json_path):
            with open(blocos_json_path, 'r', encoding='utf-8') as f:
                blocos = json.load(f)
                logging.info(f"Blocos carregados diretamente do JSON: {len(blocos)} blocos encontrados")
                
                # Adicionar título se não existir
                for nome, dados in blocos.items():
                    if 'titulo' not in dados:
                        blocos[nome]['titulo'] = nome.replace('_', ' ')
                
                return blocos
        else:
            logging.warning(f"Arquivo de blocos não encontrado: {blocos_json_path}")
            return {}
    except Exception as e:
        logging.error(f"Erro ao carregar blocos diretamente do JSON: {e}")
        return {}

# Rota para upload de imagens no editor
@app.route('/upload_image', methods=['POST'])
@login_required
def upload_image():
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "Nenhum arquivo enviado"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"success": False, "error": "Nome de arquivo vazio"}), 400
        
        if file and allowed_file(file.filename):
            # Gerar nome único para o arquivo
            filename = secure_filename(f"editor_{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}")
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # Salvar o arquivo
            file.save(file_path)
            
            # Retornar o URL da imagem
            image_url = url_for('static', filename=f"uploads/{filename}")
            return jsonify({"success": True, "url": image_url})
        else:
            return jsonify({"success": False, "error": "Formato de arquivo não permitido"}), 400
    
    except Exception as e:
        logging.error(f"Erro ao fazer upload de imagem: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

# Nova rota para visualizar logs
@app.route('/visualizar_logs')
@login_required
def visualizar_logs():
    # Verificar se o usuário é administrador
    if 'usuario_logado' not in session or session.get('superusuario') != True:
        flash('Acesso restrito a administradores.', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        log_path = os.path.join(LOGS_FOLDER, 'app.log')
        
        # Verificar se o arquivo existe
        if not os.path.exists(log_path):
            flash('Arquivo de logs não encontrado.', 'warning')
            return render_template('visualizar_logs.html', logs='Nenhum log disponível')
        
        # Ler as últimas 500 linhas do arquivo de log (para evitar carregar arquivos muito grandes)
        with open(log_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            last_lines = lines[-500:] if len(lines) > 500 else lines
            logs_content = ''.join(last_lines)
        
        # Formatação básica para exibição
        logs_content = logs_content.replace('\n', '<br>')
        
        # Destacar erros
        logs_content = logs_content.replace('ERROR', '<span style="color: red; font-weight: bold;">ERROR</span>')
        logs_content = logs_content.replace('WARNING', '<span style="color: orange; font-weight: bold;">WARNING</span>')
        logs_content = logs_content.replace('INFO', '<span style="color: green;">INFO</span>')
        
        return render_template('visualizar_logs.html', logs=logs_content)
    except Exception as e:
        flash(f'Erro ao ler logs: {str(e)}', 'danger')
        return render_template('visualizar_logs.html', logs='Erro ao ler logs')

# Rota para gerenciar ofertas
@app.route('/gerenciar_ofertas', methods=['GET', 'POST'])
@admin_required
def exibir_gerenciar_ofertas():
    # Carregar ofertas existentes
    ofertas = carregar_ofertas()
    
    # Carregar blocos para associar a ofertas
    blocos = carregar_blocos()
    
    if request.method == 'POST':
        acao = request.form.get('acao')
        
        if acao == 'adicionar':
            # Processar adição de nova oferta
            tipo_oferta = request.form.get('tipo_oferta')
            descricao = request.form.get('descricao')
            blocos_selecionados = request.form.getlist('blocos_selecionados')
            blocos_obrigatorios = request.form.getlist('blocos_obrigatorios')
            
            # Verificar se o tipo de oferta já existe
            if tipo_oferta in ofertas:
                flash(f'A oferta {tipo_oferta} já existe.', 'warning')
            else:
                # Adicionar nova oferta
                ofertas[tipo_oferta] = {
                    'descricao': descricao,
                    'blocos': {},
                    'obrigatorios': blocos_obrigatorios
                }
                
                # Processar cada bloco selecionado
                for bloco_nome in blocos_selecionados:
                    if bloco_nome in blocos:
                        ofertas[tipo_oferta]['blocos'][bloco_nome] = {
                            'titulo': blocos[bloco_nome].get('titulo', bloco_nome.replace('_', ' ')),
                            'texto': blocos[bloco_nome].get('texto', ''),
                        }
                
                # Salvar ofertas
                salvar_ofertas(ofertas)
                flash(f'Oferta {tipo_oferta} adicionada com sucesso!', 'success')
                return redirect(url_for('exibir_gerenciar_ofertas'))
                
        elif acao == 'editar':
            # Processar edição de oferta existente
            tipo_oferta = request.form.get('tipo_oferta')
            descricao = request.form.get('descricao')
            blocos_selecionados = request.form.getlist('blocos_selecionados')
            blocos_obrigatorios = request.form.getlist('blocos_obrigatorios')
            
            # Verificar se a oferta existe
            if tipo_oferta not in ofertas:
                flash(f'A oferta {tipo_oferta} não existe.', 'warning')
            else:
                # Atualizar oferta existente
                ofertas[tipo_oferta]['descricao'] = descricao
                ofertas[tipo_oferta]['obrigatorios'] = blocos_obrigatorios
                
                # Limpar e repopular blocos
                ofertas[tipo_oferta]['blocos'] = {}
                for bloco_nome in blocos_selecionados:
                    if bloco_nome in blocos:
                        ofertas[tipo_oferta]['blocos'][bloco_nome] = {
                            'titulo': blocos[bloco_nome].get('titulo', bloco_nome.replace('_', ' ')),
                            'texto': blocos[bloco_nome].get('texto', ''),
                        }
                
                # Salvar ofertas
                salvar_ofertas(ofertas)
                flash(f'Oferta {tipo_oferta} atualizada com sucesso!', 'success')
                return redirect(url_for('exibir_gerenciar_ofertas'))
                
        elif acao == 'excluir':
            # Processar exclusão de oferta
            tipo_oferta = request.form.get('tipo_oferta')
            
            # Verificar se a oferta existe
            if tipo_oferta not in ofertas:
                flash(f'A oferta {tipo_oferta} não existe.', 'warning')
            else:
                # Remover oferta
                del ofertas[tipo_oferta]
                
                # Salvar ofertas
                salvar_ofertas(ofertas)
                flash(f'Oferta {tipo_oferta} excluída com sucesso!', 'success')
                return redirect(url_for('exibir_gerenciar_ofertas'))
    
    # Renderizar template
    return render_template('gerenciar_ofertas.html', 
                          ofertas=ofertas, 
                          blocos=blocos)

# Função para salvar proposta no banco de dados - Modificada
def salvar_proposta(nome_cliente, arquivo, gerado_por, blocos_selecionados, oferta_selecionada=None):
    """
    Salva uma proposta diretamente no banco de dados e retorna seu ID.
    """
    try:
        return criar_proposta_db(nome_cliente, arquivo, gerado_por, blocos_selecionados, oferta_selecionada)
    except Exception as e:
        logger.error(f"Erro ao salvar proposta no banco: {e}")
        raise

# Função para remover rascunho - Modificada para banco de dados
def remover_rascunho(rascunho_id):
    """
    Remove um rascunho diretamente do banco de dados.
    """
    try:
        from models import Rascunho, db
        rascunho = Rascunho.query.get(rascunho_id)
        if rascunho:
            db.session.delete(rascunho)
            db.session.commit()
            logger.info(f"Rascunho removido: {rascunho_id}")
            return True
        else:
            logger.warning(f"Rascunho não encontrado: {rascunho_id}")
            return False
    except Exception as e:
        logger.error(f"Erro ao remover rascunho: {e}")
        db.session.rollback()
        return False

# Rota para API de ofertas
@app.route('/api/ofertas', methods=['GET'])
@login_required
def api_ofertas():
    try:
        ofertas = carregar_ofertas()
        return jsonify(ofertas), 200
    except Exception as e:
        app.logger.error(f"Erro ao carregar ofertas: {e}")
        return jsonify({"erro": "Erro ao carregar ofertas"}), 500

# Rota para API de oferta específica
@app.route('/api/oferta/<tipo_oferta>', methods=['GET'])
@login_required
def api_oferta(tipo_oferta):
    try:
        ofertas = carregar_ofertas()
        if tipo_oferta in ofertas:
            oferta = ofertas[tipo_oferta]
            
            # Formatar a oferta para ser compatível com o frontend
            resultado = {
                'descricao': oferta.get('descricao', ''),
                'obrigatorios': oferta.get('obrigatorios', []),
                'blocos': list(oferta.get('blocos', {}).keys())
            }
            
            app.logger.info(f"Oferta formatada para API: {resultado}")
            return jsonify(resultado), 200
        else:
            return jsonify({"erro": f"Oferta {tipo_oferta} não encontrada"}), 404
    except Exception as e:
        app.logger.error(f"Erro ao carregar oferta {tipo_oferta}: {e}")
        return jsonify({"erro": f"Erro ao carregar oferta {tipo_oferta}"}), 500

# Rota para salvar oferta via API
@app.route('/api/salvar_oferta', methods=['POST'])
@admin_required
def salvar_oferta_api():
    try:
        data = request.json
        acao = data.get('acao')
        tipo_oferta = data.get('tipo_oferta')
        descricao = data.get('descricao', '')
        blocos_selecionados = data.get('blocos_selecionados', [])
        blocos_obrigatorios = data.get('blocos_obrigatorios', [])
        
        # Carregar ofertas e blocos
        ofertas = carregar_ofertas()
        blocos = carregar_blocos()
        
        if acao == 'adicionar':
            if tipo_oferta in ofertas:
                return jsonify({"erro": f"A oferta {tipo_oferta} já existe."}), 400
                
            ofertas[tipo_oferta] = {
                'descricao': descricao,
                'blocos': {},
                'obrigatorios': blocos_obrigatorios
            }
            
            for bloco_nome in blocos_selecionados:
                if bloco_nome in blocos:
                    ofertas[tipo_oferta]['blocos'][bloco_nome] = {
                        'titulo': blocos[bloco_nome].get('titulo', bloco_nome.replace('_', ' ')),
                        'texto': blocos[bloco_nome].get('texto', ''),
                    }
            
            salvar_ofertas(ofertas)
            return jsonify({"sucesso": f"Oferta {tipo_oferta} adicionada com sucesso!"}), 201
            
        elif acao == 'editar':
            if tipo_oferta not in ofertas:
                return jsonify({"erro": f"A oferta {tipo_oferta} não existe."}), 404
                
            ofertas[tipo_oferta]['descricao'] = descricao
            ofertas[tipo_oferta]['obrigatorios'] = blocos_obrigatorios
            
            ofertas[tipo_oferta]['blocos'] = {}
            for bloco_nome in blocos_selecionados:
                if bloco_nome in blocos:
                    ofertas[tipo_oferta]['blocos'][bloco_nome] = {
                        'titulo': blocos[bloco_nome].get('titulo', bloco_nome.replace('_', ' ')),
                        'texto': blocos[bloco_nome].get('texto', ''),
                    }
            
            salvar_ofertas(ofertas)
            return jsonify({"sucesso": f"Oferta {tipo_oferta} atualizada com sucesso!"}), 200
            
        elif acao == 'excluir':
            if tipo_oferta not in ofertas:
                return jsonify({"erro": f"A oferta {tipo_oferta} não existe."}), 404
                
            del ofertas[tipo_oferta]
            salvar_ofertas(ofertas)
            return jsonify({"sucesso": f"Oferta {tipo_oferta} excluída com sucesso!"}), 200
        
        return jsonify({"erro": "Ação inválida"}), 400
    
    except Exception as e:
        app.logger.error(f"Erro ao processar oferta: {e}")
        return jsonify({"erro": f"Erro ao processar oferta: {str(e)}"}), 500

# Rota para verificar o status da API
@app.route('/api/status', methods=['GET'])
def api_status():
    try:
        # Verificar acesso ao arquivo de ofertas
        ofertas_existe = os.path.exists(OFERTAS_FILE)
        ofertas_legivel = os.access(OFERTAS_FILE, os.R_OK) if ofertas_existe else False
        
        # Verificar acesso ao arquivo de blocos
        blocos_existe = os.path.exists(BLOCOS_FILE)
        blocos_legivel = os.access(BLOCOS_FILE, os.R_OK) if blocos_existe else False
        
        # Coletar informações sobre arquivos
        status = {
            "status": "online",
            "timestamp": datetime.datetime.now().isoformat(),
            "files": {
                "ofertas": {
                    "exists": ofertas_existe,
                    "readable": ofertas_legivel,
                    "path": OFERTAS_FILE
                },
                "blocos": {
                    "exists": blocos_existe,
                    "readable": blocos_legivel,
                    "path": BLOCOS_FILE
                }
            },
            "api_routes": {
                "ofertas": "/api/ofertas",
                "oferta_especifica": "/api/oferta/<tipo_oferta>",
                "salvar_oferta": "/api/salvar_oferta"
            }
        }
        
        return jsonify(status)
    except Exception as e:
        app.logger.error(f"Erro ao verificar status da API: {e}")
        return jsonify({
            "status": "error",
            "error": str(e),
            "timestamp": datetime.datetime.now().isoformat()
        }), 500

# Função para verificar status da API periodicamente
def verificar_api_periodicamente():
    """Função para verificar o status da API periodicamente"""
    with app.app_context():
        try:
            app.logger.info("Verificando status da API...")
            # Chamar a rota de status da API
            with app.test_client() as client:
                response = client.get('/api/status')
                if response.status_code == 200:
                    app.logger.info("API respondendo normalmente.")
                else:
                    app.logger.warning(f"API respondeu com código: {response.status_code}")
        except Exception as e:
            app.logger.error(f"Erro ao verificar status da API: {str(e)}")

# Configurar agendador apenas para verificação de API (sem sincronização de JSON)
scheduler = BackgroundScheduler()
scheduler.add_job(
    func=verificar_api_periodicamente,
    trigger=IntervalTrigger(minutes=5),  # Executar a cada 5 minutos
    id='verificar_api_status',
    name='Verificar status da API',
    replace_existing=True
)

# Iniciar o agendador quando a aplicação iniciar
scheduler.start()

# Registrar função para parar o agendador quando a aplicação for encerrada
atexit.register(lambda: scheduler.shutdown())

# Função para realizar a migração inicial dos dados do JSON para o banco na inicialização
def realizar_migracao_inicial():
    """
    Realiza uma migração única dos dados dos arquivos JSON para o banco de dados
    na inicialização da aplicação, garantindo que o banco tenha todos os dados.
    """
    with app.app_context():
        try:
            app.logger.info("Realizando migração inicial de dados JSON para o banco...")
            # Verificar se já foi realizada
            from models import db, Oferta, BlocoProposta, Proposta, Rascunho
            
            # Verificar se já existem dados no banco
            ofertas_count = Oferta.query.count()
            blocos_count = BlocoProposta.query.count()
            propostas_count = Proposta.query.count()
            rascunhos_count = Rascunho.query.count()
            
            if ofertas_count == 0 or blocos_count == 0:
                app.logger.info("Banco de dados vazio ou incompleto. Realizando migração inicial...")
                
                # Importar funções de migração
                from migrate_data import migrar_perfis, migrar_usuarios, migrar_ofertas, migrar_blocos, migrar_propostas, migrar_rascunhos
                
                # Migrar dados em ordem apropriada
                migrar_perfis()
                migrar_usuarios()
                migrar_ofertas()
                migrar_blocos()
                migrar_propostas()
                migrar_rascunhos()
                
                app.logger.info("Migração inicial concluída com sucesso!")
            else:
                app.logger.info(f"Banco de dados já contém dados: {ofertas_count} ofertas, {blocos_count} blocos, {propostas_count} propostas, {rascunhos_count} rascunhos")
        except Exception as e:
            app.logger.error(f"Erro na migração inicial: {str(e)}")

# Realizar migração inicial quando a aplicação iniciar
realizar_migracao_inicial()

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080)
