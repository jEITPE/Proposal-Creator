import os
import re
import base64
import uuid
from io import BytesIO
from PIL import Image
from flask import Flask, render_template, request, redirect, url_for, send_file, abort, jsonify, session, flash
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import json
from functools import wraps
from datetime import datetime
import shutil
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask("Gerador de Propostas de Serviços Gerenciados Service IT")
app.config['IMAGE_FOLDER'] = 'static/img'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SECRET_KEY'] = 'service_it_secret_key'  # Chave para sessões
os.makedirs(app.config['IMAGE_FOLDER'], exist_ok=True)
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Arquivo para armazenar usuários
USUARIOS_FILE = os.path.join('data', 'usuarios.json')

# Função para carregar usuários
def carregar_usuarios():
    if os.path.exists(USUARIOS_FILE):
        try:
            with open(USUARIOS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {"admin": {"senha": "admin123", "tipo": "admin"}}
    else:
        # Criar arquivo de usuários com admin padrão
        usuarios = {"admin": {"senha": "admin123", "tipo": "admin"}}
        os.makedirs(os.path.dirname(USUARIOS_FILE), exist_ok=True)
        with open(USUARIOS_FILE, 'w', encoding='utf-8') as f:
            json.dump(usuarios, f, ensure_ascii=False, indent=4)
        return usuarios

# Função para salvar usuários
def salvar_usuarios(usuarios):
    os.makedirs(os.path.dirname(USUARIOS_FILE), exist_ok=True)
    with open(USUARIOS_FILE, 'w', encoding='utf-8') as f:
        json.dump(usuarios, f, ensure_ascii=False, indent=4)

# Arquivo para armazenar propostas
PROPOSTAS_FILE = os.path.join('data', 'propostas.json')

# Função para carregar propostas
def carregar_propostas():
    if os.path.exists(PROPOSTAS_FILE):
        try:
            with open(PROPOSTAS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    else:
        # Criar arquivo de propostas vazio
        propostas = {}
        os.makedirs(os.path.dirname(PROPOSTAS_FILE), exist_ok=True)
        with open(PROPOSTAS_FILE, 'w', encoding='utf-8') as f:
            json.dump(propostas, f, ensure_ascii=False, indent=4)
        return propostas

# Função para salvar propostas
def salvar_propostas(propostas):
    os.makedirs(os.path.dirname(PROPOSTAS_FILE), exist_ok=True)
    with open(PROPOSTAS_FILE, 'w', encoding='utf-8') as f:
        json.dump(propostas, f, ensure_ascii=False, indent=4)

# Arquivo para armazenar blocos
BLOCOS_FILE = os.path.join('data', 'blocos.json')

# Função para carregar blocos de conteúdo do arquivo JSON
def carregar_blocos():
    try:
        if os.path.exists(BLOCOS_FILE):
            with open(BLOCOS_FILE, "r", encoding="utf-8") as f:
                blocos = json.load(f)
                return blocos
        else:
            return {}
    except json.JSONDecodeError as e:
        print(f"Erro ao ler blocos.json: {e}")
        return {}

# Função para salvar os blocos de conteúdo no arquivo JSON
def salvar_blocos(blocos):
    try:
        os.makedirs(os.path.dirname(BLOCOS_FILE), exist_ok=True)
        with open(BLOCOS_FILE, "w", encoding="utf-8") as f:
            json.dump(blocos, f, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"Erro ao salvar blocos.json: {e}")

# Função para carregar rascunhos
def carregar_rascunhos():
    rascunhos_path = os.path.join(app.root_path, 'data', 'rascunhos.json')
    if not os.path.exists(rascunhos_path):
        return {}
    with open(rascunhos_path, 'r', encoding='utf-8') as f:
        return json.load(f)

# Função para salvar rascunhos
def salvar_rascunhos(rascunhos):
    with open(os.path.join(app.root_path, 'data', 'rascunhos.json'), 'w', encoding='utf-8') as f:
        json.dump(rascunhos, f, ensure_ascii=False, indent=4)

# Decorator para verificar se o usuário está logado
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'usuario_logado' not in session:
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

# Decorator para verificar se o usuário é admin
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'usuario_logado' not in session:
            return redirect(url_for('login', next=request.url))
        
        usuarios = carregar_usuarios()
        if session['usuario_logado'] not in usuarios or usuarios[session['usuario_logado']]['tipo'] != 'admin':
            flash('Acesso negado. Você precisa ser administrador para acessar esta página.', 'danger')
            return redirect(url_for('dashboard'))
        
        return f(*args, **kwargs)
    return decorated_function

# Rota de login
@app.route('/login', methods=['GET', 'POST'])
def login():
    # Se o usuário já estiver logado, redireciona para o dashboard
    if 'usuario_logado' in session:
        return redirect(url_for('dashboard'))
        
    error = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        usuarios = carregar_usuarios()
        if username in usuarios and usuarios[username]['senha'] == password:
            session['usuario_logado'] = username
            session['tipo_usuario'] = usuarios[username]['tipo']
            return redirect(url_for('dashboard'))
        else:
            error = 'Usuário ou senha inválidos. Tente novamente.'
    
    return render_template('login.html', error=error)

# Rota de logout
@app.route('/logout')
def logout():
    session.pop('usuario_logado', None)
    session.pop('tipo_usuario', None)
    return redirect(url_for('login'))

# Rota para o dashboard
@app.route('/dashboard')
@login_required
def dashboard():
    # Carregar propostas existentes
    propostas = carregar_propostas()
    
    # Carregar rascunhos do usuário atual
    rascunhos = carregar_rascunhos()
    rascunhos_usuario = {}
    
    usuario_atual = session.get('usuario_logado')
    for rascunho_id, rascunho in rascunhos.items():
        if rascunho.get('usuario') == usuario_atual:
            rascunhos_usuario[rascunho_id] = rascunho
    
    return render_template('dashboard.html', 
                          propostas=propostas,
                          rascunhos_usuario=rascunhos_usuario,
                          tipo_usuario=session.get('tipo_usuario', ''))

# Rota para visualizar uma proposta específica
@app.route('/visualizar_proposta/<proposta_id>')
@login_required
def visualizar_proposta(proposta_id):
    propostas = carregar_propostas()
    if proposta_id in propostas:
        proposta = propostas[proposta_id]
        
        # Carregar blocos específicos para este cliente
        blocos = carregar_blocos()
        blocos_cliente = []
        
        # Filtrar blocos específicos para este cliente
        for nome_bloco, info_bloco in blocos.items():
            if info_bloco.get('cliente_associado') == proposta['nome_cliente']:
                blocos_cliente.append(nome_bloco)
        
        return render_template('visualizar_proposta.html', 
                              proposta=proposta,
                              proposta_id=proposta_id,
                              blocos_cliente=blocos_cliente)
    else:
        flash('Proposta não encontrada.')
        return redirect(url_for('dashboard'))

# Rota para baixar uma proposta
@app.route('/baixar_proposta/<proposta_id>')
@login_required
def baixar_proposta(proposta_id):
    propostas = carregar_propostas()
    if proposta_id in propostas:
        proposta = propostas[proposta_id]
        arquivo = proposta.get('arquivo', '')
        if arquivo and os.path.exists(arquivo):
            return send_file(arquivo, as_attachment=True)
        else:
            flash('Arquivo da proposta não encontrado.')
    else:
        flash('Proposta não encontrada.')
    return redirect(url_for('dashboard'))

# Rota para excluir uma proposta
@app.route('/excluir_proposta/<proposta_id>')
@login_required
def excluir_proposta(proposta_id):
    propostas = carregar_propostas()
    if proposta_id in propostas:
        # Verificar se o usuário é admin ou o criador da proposta
        if session.get('tipo_usuario') == 'admin' or propostas[proposta_id].get('gerado_por') == session.get('usuario_logado'):
            # Remover o arquivo se existir
            arquivo = propostas[proposta_id].get('arquivo', '')
            if arquivo and os.path.exists(arquivo):
                try:
                    os.remove(arquivo)
                except:
                    pass
            # Remover a proposta do dicionário
            del propostas[proposta_id]
            # Salvar as alterações
            salvar_propostas(propostas)
            flash('Proposta excluída com sucesso.')
        else:
            flash('Você não tem permissão para excluir esta proposta.')
    else:
        flash('Proposta não encontrada.')
    return redirect(url_for('dashboard'))

# Rota para criar uma nova proposta (exibir formulário)
@app.route('/criar_proposta', methods=['GET'])
@login_required
def exibir_criar_proposta():
    # Obter o cliente da query string (se existir)
    cliente = request.args.get('cliente', '')
    
    # Verificar se hÃ¡ um rascunho para continuar
    rascunho_id = request.args.get('rascunho_id', '')
    rascunho_data = {}
    
    if rascunho_id:
        rascunhos = carregar_rascunhos()
        if rascunho_id in rascunhos and rascunhos[rascunho_id].get('usuario') == session.get('usuario_logado'):
            rascunho_data = rascunhos[rascunho_id]
            cliente = rascunho_data.get('nome_cliente', cliente)
    
    # Limpar qualquer seleÃ§Ã£o anterior de blocos (para reiniciar a pÃ¡gina)
    if 'blocos_selecionados' in session and not rascunho_id:
        session.pop('blocos_selecionados')
    
    # Carregar blocos de texto
    blocos = carregar_blocos()
    
    # Filtrar blocos: mostrar apenas blocos obrigatÃ³rios e blocos especÃ­ficos para este cliente
    usuario_atual = session.get('usuario_logado')
    blocos_filtrados = {}
    
    for nome_bloco, info_bloco in blocos.items():
        # Incluir blocos obrigatÃ³rios
        if info_bloco.get('obrigatorio', False):
            blocos_filtrados[nome_bloco] = info_bloco
        # Incluir blocos sem cliente associado (blocos gerais)
        elif not info_bloco.get('cliente_associado'):
            blocos_filtrados[nome_bloco] = info_bloco
        # Incluir blocos especÃ­ficos para este cliente
        elif info_bloco.get('cliente_associado') == cliente:
            blocos_filtrados[nome_bloco] = info_bloco
        # Incluir blocos criados pelo usuÃ¡rio atual
        elif info_bloco.get('criado_por') == usuario_atual:
            blocos_filtrados[nome_bloco] = info_bloco
    
    # Obter lista de modelos disponÃ­veis
    modelos = obter_modelos_disponiveis()
    
    # Obter tipo de usuÃ¡rio
    is_admin = session.get('tipo_usuario') == 'admin'
    
    return render_template('criar_proposta.html', 
                          blocos=blocos_filtrados, 
                          modelos=modelos,
                          is_admin=is_admin,
                          cliente=cliente,
                          rascunho=rascunho_data,
                          rascunho_id=rascunho_id)

# Rota para processar a criaÃ§Ã£o de uma proposta
@app.route('/criar_proposta', methods=['POST'])
@login_required
def criar_proposta():
    nome_cliente = request.form.get('nome_cliente', '')
    logo_cliente = request.form.get('logo_cliente', '')
    modelo_proposta = request.form.get('modelo_proposta', '')
    blocos_selecionados = request.form.getlist('blocos')
    rascunho_id = request.form.get('rascunho_id', '')
    
    # Verificar se Ã© para salvar como rascunho
    if 'salvar_rascunho' in request.form:
        return salvar_como_rascunho(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados, rascunho_id)
    
    # Se nÃ£o houver modelo selecionado, use o primeiro disponÃ­vel
    if not modelo_proposta:
        modelos = obter_modelos_disponiveis()
        if modelos:
            modelo_proposta = modelos[0]
        else:
            flash('Nenhum modelo de proposta disponÃ­vel.')
            return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
    
    # Processar upload de arquivo de logo
    logo_path = None
    if 'logo_file' in request.files and request.files['logo_file'].filename:
        logo_file = request.files['logo_file']
        # Verificar extensÃ£o
        if logo_file and '.' in logo_file.filename and logo_file.filename.rsplit('.', 1)[1].lower() in ['png', 'jpg', 'jpeg', 'gif']:
            # Gerar nome Ãºnico para o arquivo
            filename = secure_filename(f"{nome_cliente.replace(' ', '_')}_{uuid.uuid4()}.{logo_file.filename.rsplit('.', 1)[1].lower()}")
            logo_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            logo_file.save(logo_path)
        else:
            flash('Formato de arquivo nÃ£o suportado. Use PNG, JPG ou GIF.')
            return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
    elif logo_cliente:
        # Se foi fornecida uma URL, baixar a imagem
        try:
            import requests
            from urllib.parse import urlparse
            
            response = requests.get(logo_cliente, stream=True)
            if response.status_code == 200:
                # Extrair extensÃ£o da URL ou usar .png como padrÃ£o
                parsed_url = urlparse(logo_cliente)
                path = parsed_url.path
                ext = os.path.splitext(path)[1]
                if not ext:
                    ext = '.png'
                
                # Gerar nome Ãºnico para o arquivo
                filename = secure_filename(f"{nome_cliente.replace(' ', '_')}_{uuid.uuid4()}{ext}")
                logo_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                
                with open(logo_path, 'wb') as f:
                    for chunk in response.iter_content(1024):
                        f.write(chunk)
            else:
                flash('NÃ£o foi possÃ­vel baixar a imagem da URL fornecida.')
                return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
        except Exception as e:
            flash(f'Erro ao processar a URL da imagem: {str(e)}')
            return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
    
    # Validar dados
    if not nome_cliente or not modelo_proposta:
        flash('Por favor, preencha todos os campos obrigatÃ³rios.')
        return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))
    
    # Carregar blocos de texto e verificar blocos obrigatÃ³rios
    blocos_texto = carregar_blocos()
    blocos_obrigatorios = [bloco for bloco, info in blocos_texto.items() if info.get('obrigatorio', False)]
    
    # Adicionar blocos obrigatÃ³rios se nÃ£o estiverem na lista
    for bloco in blocos_obrigatorios:
        if bloco not in blocos_selecionados:
            blocos_selecionados.append(bloco)
    
    # Gerar a proposta
    try:
        arquivo_gerado = gerar_proposta(nome_cliente, logo_path, modelo_proposta, blocos_selecionados)
        
        # Criar registro da proposta
        proposta_id = str(uuid.uuid4())
        propostas = carregar_propostas()
        
        propostas[proposta_id] = {
            'nome_cliente': nome_cliente,
            'data_geracao': datetime.now().strftime('%d/%m/%Y %H:%M'),
            'gerado_por': session.get('usuario_logado', 'Desconhecido'),
            'arquivo': arquivo_gerado,
            'blocos_utilizados': blocos_selecionados
        }
        
        salvar_propostas(propostas)
        flash('Proposta gerada com sucesso!')
        return redirect(url_for('visualizar_proposta', proposta_id=proposta_id))
    except Exception as e:
        flash(f'Erro ao gerar proposta: {str(e)}')
        return redirect(url_for('exibir_criar_proposta', cliente=nome_cliente))

# FunÃ§Ã£o para obter modelos disponÃ­veis
def obter_modelos_disponiveis():
    modelos_dir = os.path.join(app.root_path, 'modelos')
    if not os.path.exists(modelos_dir):
        return []
    
    modelos = []
    for arquivo in os.listdir(modelos_dir):
        if arquivo.endswith('.docx') and not arquivo.startswith('~$'):
            modelos.append(arquivo)
    
    return modelos

# FunÃ§Ã£o para gerar proposta
def gerar_proposta(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados):
    # Carregar o modelo
    modelo_path = os.path.join(app.root_path, 'modelos', modelo_proposta)
    if not os.path.exists(modelo_path):
        raise Exception(f"Modelo '{modelo_proposta}' nÃ£o encontrado.")
    
    # Carregar blocos de texto
    blocos_texto = carregar_blocos()
    
    # Criar diretÃ³rio para propostas geradas se nÃ£o existir
    output_dir = os.path.join(app.root_path, 'propostas_geradas')
    os.makedirs(output_dir, exist_ok=True)
    
    # Nome do arquivo de saÃ­da
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_file = os.path.join(output_dir, f"Proposta_{nome_cliente.replace(' ', '_')}_{timestamp}.docx")
    
    # Copiar o modelo para o arquivo de saÃ­da
    shutil.copy2(modelo_path, output_file)
    
    # Abrir o documento
    doc = Document(output_file)
    
    # Substituir a capa
    substituir_capa(doc, nome_cliente, logo_cliente)
    
    # Encontrar o marcador onde inserir os blocos
    marcador_encontrado = False
    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{blocos_conteudo}}" in paragrafo.text or "blocos_conteudo" in paragrafo.text:
            marcador_encontrado = True
            # Remover o parÃ¡grafo do marcador
            p = paragrafo._p
            p.getparent().remove(p)
            
            # Adicionar os blocos selecionados
            for bloco in blocos_selecionados:
                if bloco in blocos_texto:
                    conteudo = blocos_texto[bloco].get('texto', '')
                    # Adicionar tÃ­tulo do bloco
                    titulo = doc.add_paragraph()
                    titulo.style = 'Heading 2'
                    titulo.add_run(bloco.replace('_', ' ').capitalize())
                    
                    # Adicionar conteÃºdo do bloco
                    doc.add_paragraph(conteudo)
            break
    
    # Se nÃ£o encontrou o marcador, adicionar os blocos no final
    if not marcador_encontrado:
        doc.add_paragraph("ConteÃºdo da Proposta", style='Heading 1')
        for bloco in blocos_selecionados:
            if bloco in blocos_texto:
                conteudo = blocos_texto[bloco].get('texto', '')
                # Adicionar tÃ­tulo do bloco
                titulo = doc.add_paragraph()
                titulo.style = 'Heading 2'
                titulo.add_run(bloco.replace('_', ' ').capitalize())
                
                # Adicionar conteÃºdo do bloco
                doc.add_paragraph(conteudo)
    
    # Salvar o documento
    doc.save(output_file)
    
    return output_file

# FunÃ§Ã£o para substituir o marcador {{nome_cliente}} e {{logo_cliente}} na capa do documento
def substituir_capa(document, nome_cliente, logo_cliente_path):
    # Substituir nos parÃ¡grafos principais
    for paragraph in document.paragraphs:
        # Substituir texto do cliente em cada run
        for run in paragraph.runs:
            if ("{{nome_cliente}}" in run.text or "{{NOME_CLIENTE}}" in run.text or "nome_cliente" in run.text.lower()):
                run.text = run.text.replace("{{nome_cliente}}", nome_cliente)\
                                   .replace("{{NOME_CLIENTE}}", nome_cliente)\
                                   .replace("NOME_CLIENTE", nome_cliente)
        # Inserir logotipo se o marcador for encontrado
        if ( ("{{logo_cliente}}" in paragraph.text or "logo_cliente" in paragraph.text) and logo_cliente_path):
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(logo_cliente_path, width=Inches(1.5))
    
    # Substituir nos cabeÃ§alhos e rodapÃ©s
    for section in document.sections:
        # CabeÃ§alhos
        for header in [section.header]:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    if ("{{nome_cliente}}" in run.text or "{{NOME_CLIENTE}}" in run.text or "nome_cliente" in run.text.lower()):
                        run.text = run.text.replace("{{nome_cliente}}", nome_cliente)\
                                           .replace("{{NOME_CLIENTE}}", nome_cliente)\
                                           .replace("NOME_CLIENTE", nome_cliente)
        # RodapÃ©s
        for footer in [section.footer]:
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    if ("{{nome_cliente}}" in run.text or "{{NOME_CLIENTE}}" in run.text or "nome_cliente" in run.text.lower()):
                        run.text = run.text.replace("{{nome_cliente}}", nome_cliente)\
                                           .replace("{{NOME_CLIENTE}}", nome_cliente)\
                                           .replace("NOME_CLIENTE", nome_cliente)
    
    # Substituir dentro das tabelas
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if ("{{nome_cliente}}" in run.text or "{{NOME_CLIENTE}}" in run.text or "nome_cliente" in run.text.lower()):
                            run.text = run.text.replace("{{nome_cliente}}", nome_cliente)\
                                               .replace("{{NOME_CLIENTE}}", nome_cliente)\
                                               .replace("NOME_CLIENTE", nome_cliente)
                    if ( ("{{logo_cliente}}" in paragraph.text or "logo_cliente" in paragraph.text) and logo_cliente_path):
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.add_picture(logo_cliente_path, width=Inches(1.5))

# Rota padrÃ£o - redireciona para login se nÃ£o estiver autenticado
@app.route('/')
@login_required
def index():
    return redirect(url_for('dashboard'))

# Rota para a pÃ¡gina inicial apÃ³s o login
@app.route('/home')
@login_required
def home():
    return redirect(url_for('index'))

# Rota para editar e salvar blocos
@app.route('/editar_blocos')
@admin_required
def editar_blocos():
    blocos_de_texto = carregar_blocos()
    is_admin = session.get('tipo_usuario') == 'admin'
    return render_template('editar_blocos.html', blocos=blocos_de_texto, is_admin=is_admin)

# Rota para salvar as ediÃ§Ãµes dos blocos no JSON
@app.route('/salvar_blocos', methods=['POST'])
@admin_required
def salvar_edicao_blocos():
    blocos_de_texto = carregar_blocos()
    is_admin = session.get('tipo_usuario') == 'admin'
    
    # Verificar se o usuÃ¡rio Ã© admin para editar blocos obrigatÃ³rios
    for bloco in blocos_de_texto.keys():
        if not is_admin and blocos_de_texto[bloco].get('obrigatorio', False):
            continue  # SDMs nÃ£o podem editar blocos obrigatÃ³rios
            
        novo_texto = request.form.get(f"{bloco}_texto")
        if novo_texto:
            # Substituir imagens embutidas por imagens salvas corretamente
            novo_texto = salvar_imagens_e_substituir(novo_texto, bloco)
            blocos_de_texto[bloco]['texto'] = novo_texto
    
    salvar_blocos(blocos_de_texto)
    return redirect(url_for('dashboard'))

# Rota para excluir um bloco de conteÃºdo
@app.route('/excluir_bloco/<nome_bloco>', methods=['GET'])
@login_required
def excluir_bloco(nome_bloco):
    try:
        blocos_de_texto = carregar_blocos()
        is_admin = session.get('tipo_usuario') == 'admin'
        
        # Verificar se o bloco existe
        if nome_bloco not in blocos_de_texto:
            flash('Bloco nÃ£o encontrado.', 'danger')
            return redirect(url_for('editar_blocos'))
        
        # Verificar se o bloco Ã© obrigatÃ³rio
        if blocos_de_texto[nome_bloco].get('obrigatorio', False):
            if not is_admin:
                flash('VocÃª nÃ£o tem permissÃ£o para excluir blocos obrigatÃ³rios.', 'danger')
                return redirect(url_for('dashboard'))
            else:
                # Confirmar exclusÃ£o de bloco obrigatÃ³rio por admin
                flash('AtenÃ§Ã£o: VocÃª estÃ¡ excluindo um bloco obrigatÃ³rio.', 'warning')
        
        # Excluir o bloco
        del blocos_de_texto[nome_bloco]
        salvar_blocos(blocos_de_texto)
        flash(f'Bloco "{nome_bloco.replace("_", " ")}" excluÃ­do com sucesso!', 'success')
    except Exception as e:
        flash(f'Erro ao excluir bloco: {str(e)}', 'danger')
    
    if is_admin:
        return redirect(url_for('editar_blocos'))
    else:
        return redirect(url_for('dashboard'))

# Rota para adicionar um novo bloco de conteÃºdo
@app.route('/adicionar_bloco', methods=['GET', 'POST'])
@login_required
def adicionar_bloco():
    is_admin = session.get('tipo_usuario') == 'admin'
    cliente = request.args.get('cliente', '')
    
    if request.method == 'POST':
        nome_bloco = request.form['nome_bloco'].replace(' ', '_')
        texto_bloco = request.form['texto_bloco']
        obrigatorio = request.form.get('obrigatorio') == 'on'
        cliente_associado = request.form.get('cliente_associado', '')
        
        # Apenas admins podem marcar blocos como obrigatÃ³rios
        if not is_admin:
            obrigatorio = False
        
        blocos_de_texto = carregar_blocos()
        
        # Verificar se o bloco jÃ¡ existe
        if nome_bloco in blocos_de_texto:
            flash(f'JÃ¡ existe um bloco com o nome "{nome_bloco.replace("_", " ")}"!', 'danger')
            return render_template('adicionar_bloco.html', is_admin=is_admin, cliente=cliente_associado)
        
        # Salvar imagens embutidas no texto HTML
        texto_processado = salvar_imagens_e_substituir(texto_bloco, nome_bloco)
        
        # Adicionar o novo bloco
        blocos_de_texto[nome_bloco] = {
            'texto': texto_processado,
            'imagem': None,
            'obrigatorio': obrigatorio,
            'criado_por': session.get('usuario_logado'),
            'data_criacao': datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            'cliente_associado': cliente_associado  # Novo campo para associar o bloco a um cliente
        }
        
        salvar_blocos(blocos_de_texto)
        flash(f'Bloco "{nome_bloco.replace("_", " ")}" adicionado com sucesso!', 'success')
        
        # Redirecionar para a pÃ¡gina de criaÃ§Ã£o de proposta com o cliente preenchido
        if cliente_associado:
            return redirect(url_for('exibir_criar_proposta', cliente=cliente_associado))
        else:
            return redirect(url_for('exibir_criar_proposta'))
    
    return render_template('adicionar_bloco.html', is_admin=is_admin, cliente=cliente)

# FunÃ§Ã£o para processar imagens e salvar com caminho correto no HTML
def salvar_imagens_e_substituir(texto_html, bloco_nome):
    img_counter = 1

    def salvar_imagem(match):
        nonlocal img_counter
        img_data = match.group(1)
        if img_data.startswith("data:image/"):
            img_data = img_data.split(",")[1]
            img_bytes = base64.b64decode(img_data)
            img_nome = f"{bloco_nome}_img_{img_counter}.png"
            img_path = os.path.join(app.config['IMAGE_FOLDER'], img_nome)
            with Image.open(BytesIO(img_bytes)) as img:
                img.save(img_path, format="PNG")
            img_counter += 1
            return f'<img src="/{img_path}">'
        return match.group(0)

    return re.sub(r'<img src="([^"]+)"', salvar_imagem, texto_html)

# Rota para gerenciar usuÃ¡rios
@app.route('/gerenciar_usuarios', methods=['GET', 'POST'])
@admin_required
def gerenciar_usuarios():
    if request.method == 'POST':
        acao = request.form.get('acao')
        
        if acao == 'adicionar':
            # Adicionar novo usuÃ¡rio
            novo_usuario = request.form.get('novo_usuario')
            nova_senha = request.form.get('nova_senha')
            tipo_usuario = request.form.get('tipo_usuario')
            
            if not novo_usuario or not nova_senha or not tipo_usuario:
                flash('Todos os campos sÃ£o obrigatÃ³rios.')
                return redirect(url_for('gerenciar_usuarios'))
            
            if tipo_usuario not in ['admin', 'sdm']:
                flash('Tipo de usuÃ¡rio invÃ¡lido.')
                return redirect(url_for('gerenciar_usuarios'))
            
            usuarios = carregar_usuarios()
            
            if novo_usuario in usuarios:
                flash('Nome de usuÃ¡rio jÃ¡ existe.')
                return redirect(url_for('gerenciar_usuarios'))
            
            usuarios[novo_usuario] = {
                'senha': nova_senha,
                'tipo': tipo_usuario
            }
            
            salvar_usuarios(usuarios)
            flash('UsuÃ¡rio adicionado com sucesso.')
            
        elif acao == 'remover':
            # Remover usuÃ¡rio existente
            usuario_remover = request.form.get('usuario_remover')
            
            if usuario_remover == 'admin':
                flash('NÃ£o Ã© possÃ­vel remover o usuÃ¡rio admin principal.')
                return redirect(url_for('gerenciar_usuarios'))
            
            usuarios = carregar_usuarios()
            
            if usuario_remover in usuarios:
                del usuarios[usuario_remover]
                salvar_usuarios(usuarios)
                flash('UsuÃ¡rio removido com sucesso.')
            else:
                flash('UsuÃ¡rio nÃ£o encontrado.')
    
    usuarios = carregar_usuarios()
    return render_template('gerenciar_usuarios.html', usuarios=usuarios)

# Rota para adicionar usuÃ¡rio (mantida para compatibilidade)
@app.route('/adicionar_usuario', methods=['POST'])
@admin_required
def adicionar_usuario():
    username = request.form.get('novo_usuario')
    password = request.form.get('nova_senha')
    tipo = request.form.get('tipo_usuario')
    
    if not username or not password or not tipo:
        flash('Todos os campos sÃ£o obrigatÃ³rios.')
        return redirect(url_for('gerenciar_usuarios'))
    
    if tipo not in ['admin', 'sdm']:
        flash('Tipo de usuÃ¡rio invÃ¡lido.')
        return redirect(url_for('gerenciar_usuarios'))
    
    usuarios = carregar_usuarios()
    
    if username in usuarios:
        flash('Nome de usuÃ¡rio jÃ¡ existe.')
        return redirect(url_for('gerenciar_usuarios'))
    
    usuarios[username] = {
        'senha': password,
        'tipo': tipo
    }
    
    salvar_usuarios(usuarios)
    flash('UsuÃ¡rio adicionado com sucesso.')
    return redirect(url_for('gerenciar_usuarios'))

# FunÃ§Ã£o para salvar proposta como rascunho
def salvar_como_rascunho(nome_cliente, logo_cliente, modelo_proposta, blocos_selecionados, rascunho_id=None):
    # Carregar rascunhos existentes
    rascunhos = carregar_rascunhos()
    
    # Se nÃ£o tiver um ID de rascunho, crie um novo
    if not rascunho_id:
        rascunho_id = str(uuid.uuid4())
    
    # Criar ou atualizar o rascunho
    rascunhos[rascunho_id] = {
        'nome_cliente': nome_cliente,
        'logo_cliente': logo_cliente,
        'modelo_proposta': modelo_proposta,
        'blocos_selecionados': blocos_selecionados,
        'usuario': session.get('usuario_logado'),
        'data_atualizacao': datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    }
    
    # Salvar os rascunhos
    salvar_rascunhos(rascunhos)
    
    flash('Rascunho salvo com sucesso!')
    return redirect(url_for('dashboard'))

@app.route('/excluir_rascunho/<rascunho_id>')
@login_required
def excluir_rascunho(rascunho_id):
    # Carregar rascunhos existentes
    rascunhos = carregar_rascunhos()
    
    # Verificar se o rascunho existe
    if rascunho_id not in rascunhos:
        flash('Rascunho não encontrado.')
        return redirect(url_for('dashboard'))
    
    # Verificar se o usuário tem permissão para excluir o rascunho
    if rascunhos[rascunho_id]['usuario'] != session.get('usuario_logado') and session.get('tipo_usuario') != 'admin':
        flash('Você não tem permissão para excluir este rascunho.')
        return redirect(url_for('dashboard'))
    
    # Excluir o rascunho
    del rascunhos[rascunho_id]
    
    # Salvar os rascunhos atualizados
    salvar_rascunhos(rascunhos)
    
    flash('Rascunho excluído com sucesso!')
    return redirect(url_for('dashboard'))

if __name__ == '__main__':
    app.run(debug=True)
